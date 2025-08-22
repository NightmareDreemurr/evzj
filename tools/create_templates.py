"""
Enhanced DOCX template creator with Chinese font support.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX


def create_assignment_template():
    """Create enhanced assignment batch template."""
    doc = Document()
    
    # Set default font for Chinese text
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'  # 宋体 - good for Chinese
    font.size = Pt(12)
    
    # Title with center alignment
    title = doc.add_heading('{{ assignment.title }} - 作业批量报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'SimHei'  # 黑体 for headers
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Assignment metadata section
    doc.add_heading('作业信息', level=1)
    info_para = doc.add_paragraph()
    info_para.add_run('班级：').bold = True
    info_para.add_run('{{ assignment.classroom.name }}\n')
    info_para.add_run('教师：').bold = True  
    info_para.add_run('{{ assignment.teacher.name }}\n')
    info_para.add_run('生成时间：').bold = True
    info_para.add_run('{{ now.strftime("%Y年%m月%d日 %H:%M:%S") }}')
    
    # Set font for info paragraph
    for run in info_para.runs:
        run.font.name = 'SimSun'
        run.font.size = Pt(12)
    
    # Students section header
    doc.add_heading('学生作文评估报告', level=1)
    
    # Instructions for template processing
    instruction_para = doc.add_paragraph()
    instruction_para.add_run('{% for student in students %}').font.color.rgb = RGBColor(255, 255, 255)  # White text (hidden)
    
    # Student section template
    student_header = doc.add_heading('{{ student.meta.student }}', level=2)
    student_header_run = student_header.runs[0]
    student_header_run.font.name = 'SimHei'
    student_header_run.font.size = Pt(16)
    student_header_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
    
    # Student basic info
    student_info = doc.add_paragraph()
    student_info.add_run('题目：').bold = True
    student_info.add_run('{{ student.meta.topic }}\n')
    student_info.add_run('字数：').bold = True
    student_info.add_run('{{ student.meta.words }}字\n')
    student_info.add_run('评估日期：').bold = True
    student_info.add_run('{{ student.meta.date }}')
    
    # Set font for student info
    for run in student_info.runs:
        run.font.name = 'SimSun'
        run.font.size = Pt(11)
    
    # Scores section
    scores_header = doc.add_heading('评分结果', level=3)
    scores_header.runs[0].font.name = 'SimHei'
    scores_header.runs[0].font.size = Pt(14)
    
    # Total score
    total_score_para = doc.add_paragraph()
    total_score_para.add_run('总分：').bold = True
    total_score_para.add_run('{{ student.scores.total }}')
    total_score_para.add_run(' 分')
    
    # Individual rubric scores
    rubrics_para = doc.add_paragraph()
    rubrics_para.add_run('{% for r in student.scores.rubrics %}').font.color.rgb = RGBColor(255, 255, 255)
    rubrics_para.add_run('{{ r.name }}：{{ r.score }}/{{ r.max }} ')
    rubrics_para.add_run('{% endfor %}').font.color.rgb = RGBColor(255, 255, 255)
    
    # Set font for scores
    for para in [total_score_para, rubrics_para]:
        for run in para.runs:
            if run.font.color.rgb != RGBColor(255, 255, 255):  # Skip hidden template code
                run.font.name = 'SimSun'
                run.font.size = Pt(11)
    
    # Original text section
    original_header = doc.add_heading('原文内容', level=3)
    original_header.runs[0].font.name = 'SimHei'
    original_header.runs[0].font.size = Pt(14)
    
    original_para = doc.add_paragraph()
    original_para.add_run('{{ student.text.original or "原文不可用" }}')
    original_para.runs[0].font.name = 'SimSun'
    original_para.runs[0].font.size = Pt(11)
    original_para.style = doc.styles['Normal']
    
    # AI feedback section
    feedback_header = doc.add_heading('AI评语与建议', level=3)
    feedback_header.runs[0].font.name = 'SimHei'
    feedback_header.runs[0].font.size = Pt(14)
    
    feedback_para = doc.add_paragraph()
    feedback_para.add_run('{{ student.diagnosis.comment or "暂无评语" }}')
    feedback_para.runs[0].font.name = 'SimSun'
    feedback_para.runs[0].font.size = Pt(11)
    
    # Page break between students
    pagebreak_para = doc.add_paragraph()
    pagebreak_para.add_run('{% if not loop.last %}').font.color.rgb = RGBColor(255, 255, 255)
    
    # Add actual page break placeholder
    doc.add_page_break()
    
    pagebreak_end = doc.add_paragraph()
    pagebreak_end.add_run('{% endif %}').font.color.rgb = RGBColor(255, 255, 255)
    
    # End of student loop
    end_loop = doc.add_paragraph()
    end_loop.add_run('{% endfor %}').font.color.rgb = RGBColor(255, 255, 255)
    
    # Footer section
    doc.add_paragraph('\n' + '='*50)
    footer = doc.add_paragraph()
    footer.add_run('报告生成：e文智教系统 - ').font.name = 'SimSun'
    footer.add_run('{{ now.strftime("%Y年%m月%d日") }}').font.name = 'SimSun'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


if __name__ == '__main__':
    doc = create_assignment_template()
    doc.save('templates/word/assignment_compiled.docx')
    print('Enhanced assignment template created with Chinese font support')