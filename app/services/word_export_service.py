import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from PIL import Image
from flask import current_app

from app.models import Essay, EssayAssignment, GradingStandard, Dimension


class WordExportService:
    """Word文档导出服务"""
    
    def __init__(self):
        self.document = None
        
    def export_assignment_essays(self, assignment_id, output_path=None):
        """
        导出指定作业的所有学生作文到Word文档
        
        Args:
            assignment_id: 作业ID
            output_path: 输出文件路径，如果为None则自动生成
            
        Returns:
            str: 生成的Word文档路径
        """
        try:
            # 获取作业信息
            assignment = EssayAssignment.query.get_or_404(assignment_id)
            
            # 获取该作业的所有已评分作文
            from app.models import Enrollment, User, StudentProfile
            essays = Essay.query.filter_by(
                assignment_id=assignment_id,
                status='graded'
            ).join(Enrollment).join(StudentProfile).join(User).order_by(
                Enrollment.student_number,
                User.full_name
            ).all()
            
            if not essays:
                raise ValueError("没有找到已评分的作文")
            
            # 创建Word文档
            self.document = Document()
            
            # 设置文档标题
            self._add_document_title(assignment)
            
            # 为每个学生添加作文内容
            for i, essay in enumerate(essays):
                self._add_essay_content(essay, assignment)
                
                # 如果不是最后一个学生，添加分页符
                if i < len(essays) - 1:
                    self._add_page_break()
            
            # 生成输出路径
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"作业导出_{assignment.title}_{timestamp}.docx"
                # 确保导出目录存在
                export_dir = os.path.join(current_app.instance_path, 'exports')
                os.makedirs(export_dir, exist_ok=True)
                output_path = os.path.join(export_dir, filename)
            
            # 保存文档
            self.document.save(output_path)
            
            return output_path
            
        except Exception as e:
            current_app.logger.error(f"Word导出失败: {str(e)}")
            raise
    
    def _add_document_title(self, assignment):
        """添加文档标题"""
        # 主标题
        title = self.document.add_heading(f'作业：{assignment.title}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 作业信息
        info_para = self.document.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_text = f"教师：{assignment.teacher.user.full_name}\n"
        info_text += f"发布时间：{assignment.created_at.strftime('%Y年%m月%d日')}\n"
        if assignment.due_date:
            info_text += f"截止时间：{assignment.due_date.strftime('%Y年%m月%d日')}\n"
        info_text += f"评分标准：{assignment.grading_standard.title}\n"
        info_text += f"导出时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}"
        
        info_para.add_run(info_text)
        
        # 添加分隔线
        self.document.add_paragraph('=' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def _add_essay_content(self, essay, assignment):
        """添加单个学生的作文内容"""
        # 学生信息标题
        student_name = essay.enrollment.student.user.full_name
        student_number = essay.enrollment.student_number or "无学号"
        classroom_name = essay.enrollment.classroom.class_name
        
        student_heading = self.document.add_heading(
            f'{student_name} ({student_number}) - {classroom_name}', 1
        )
        student_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 添加作文图片（如果存在）
        if essay.original_image_path and os.path.exists(essay.original_image_path):
            self._add_essay_image(essay.original_image_path)
        
        # 添加作文内容
        self._add_essay_text(essay)
        
        # 添加评分信息
        self._add_scoring_info(essay, assignment.grading_standard)
        
        # 添加评语
        self._add_feedback(essay)
        
    def _add_essay_image(self, image_path):
        """添加作文图片"""
        try:
            # 检查图片文件是否存在
            if not os.path.exists(image_path):
                self.document.add_paragraph("[作文图片文件不存在]")
                return
                
            # 获取图片尺寸并调整
            with Image.open(image_path) as img:
                width, height = img.size
                
            # 设置最大宽度为6英寸，保持宽高比
            max_width = Inches(6)
            if width > height:
                # 横向图片
                img_width = max_width
                img_height = max_width * height / width
            else:
                # 纵向图片
                img_height = Inches(8)  # 最大高度8英寸
                img_width = img_height * width / height
                if img_width > max_width:
                    img_width = max_width
                    img_height = max_width * height / width
            
            # 添加图片标题
            self.document.add_paragraph("作文原图：").runs[0].bold = True
            
            # 添加图片
            para = self.document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0] if para.runs else para.add_run()
            run.add_picture(image_path, width=img_width, height=img_height)
            
        except Exception as e:
            current_app.logger.error(f"添加图片失败: {str(e)}")
            self.document.add_paragraph(f"[图片加载失败: {str(e)}]")
    
    def _add_essay_text(self, essay):
        """添加作文文本内容"""
        # 作文内容标题
        self.document.add_paragraph("作文内容：").runs[0].bold = True
        
        # 使用最终的作文内容（可能经过教师校正）
        content = essay.teacher_corrected_text or essay.content or essay.original_ocr_text
        
        if content:
            content_para = self.document.add_paragraph(content)
            # 设置内容段落格式
            content_para.paragraph_format.line_spacing = 1.5
            content_para.paragraph_format.space_after = Pt(12)
        else:
            self.document.add_paragraph("[无作文内容]")
    
    def _add_scoring_info(self, essay, grading_standard):
        """添加评分信息"""
        self.document.add_paragraph("评分详情：").runs[0].bold = True
        
        # 获取最终评分（教师调整后的分数或AI分数）
        final_scores = essay.teacher_score or essay.ai_score
        
        if not final_scores:
            self.document.add_paragraph("[暂无评分]")
            return
            
        # 解析评分数据
        if isinstance(final_scores, str):
            try:
                final_scores = json.loads(final_scores)
            except:
                self.document.add_paragraph("[评分数据格式错误]")
                return
        
        # 创建评分表格
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '评分维度'
        header_cells[1].text = '得分'
        header_cells[2].text = '满分'
        
        # 设置表头格式
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加各维度评分
        total_score = 0
        total_max_score = 0
        
        for dimension in grading_standard.dimensions:
            row_cells = table.add_row().cells
            row_cells[0].text = dimension.name
            
            # 获取该维度的得分
            dimension_score = final_scores.get('scores', {}).get(dimension.name, 0)
            row_cells[1].text = str(dimension_score)
            row_cells[2].text = str(dimension.max_score)
            
            # 居中对齐
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            total_score += float(dimension_score)
            total_max_score += dimension.max_score
        
        # 添加总分行
        total_row_cells = table.add_row().cells
        total_row_cells[0].text = '总分'
        total_row_cells[1].text = f'{total_score:.1f}'
        total_row_cells[2].text = str(total_max_score)
        
        # 设置总分行格式
        for cell in total_row_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_feedback(self, essay):
        """添加评语信息"""
        self.document.add_paragraph("评语建议：").runs[0].bold = True
        
        # 获取评语数据
        ai_feedback = essay.ai_score
        teacher_overrides = essay.teacher_feedback_overrides
        
        if isinstance(ai_feedback, str):
            try:
                ai_feedback = json.loads(ai_feedback)
            except:
                ai_feedback = {}
        
        if isinstance(teacher_overrides, str):
            try:
                teacher_overrides = json.loads(teacher_overrides)
            except:
                teacher_overrides = {}
        
        # 获取各维度的评语
        feedback_data = ai_feedback.get('feedback', {})
        
        if not feedback_data:
            self.document.add_paragraph("[暂无评语]")
            return
        
        # 添加各维度评语
        for dimension_name, feedback_text in feedback_data.items():
            # 检查是否有教师覆写的评语
            final_feedback = teacher_overrides.get(dimension_name, feedback_text)
            
            if final_feedback:
                # 维度名称
                dimension_para = self.document.add_paragraph()
                dimension_run = dimension_para.add_run(f"{dimension_name}：")
                dimension_run.bold = True
                
                # 评语内容
                feedback_para = self.document.add_paragraph(final_feedback)
                feedback_para.paragraph_format.left_indent = Inches(0.5)
                feedback_para.paragraph_format.space_after = Pt(6)
    
    def _add_page_break(self):
        """添加分页符"""
        self.document.add_page_break()