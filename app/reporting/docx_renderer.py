"""
DOCX rendering module for evaluation reports.
"""
import os
import tempfile
import logging
from typing import Dict, Any
from pathlib import Path

try:
    from docxtpl import DocxTemplate
    DOCXTPL_AVAILABLE = True
except ImportError:
    DOCXTPL_AVAILABLE = False

from docx import Document
from docx.shared import Inches

from app.schemas.evaluation import EvaluationResult, to_context

logger = logging.getLogger(__name__)


def _sanitize_filename(filename: str) -> str:
    """
    Sanitize filename to be safe for all operating systems.
    
    Removes or replaces characters that are invalid on Windows:
    < > : " | ? * \\ /
    
    Args:
        filename: Original filename string
        
    Returns:
        Sanitized filename safe for all OS
    """
    # Characters that are invalid on Windows
    invalid_chars = '<>:"|?*\\/'
    
    # Replace invalid characters with underscores
    sanitized = filename
    for char in invalid_chars:
        sanitized = sanitized.replace(char, '_')
    
    # Replace spaces with underscores for consistency
    sanitized = sanitized.replace(' ', '_')
    
    # Remove multiple consecutive underscores
    while '__' in sanitized:
        sanitized = sanitized.replace('__', '_')
    
    # Remove leading/trailing underscores
    sanitized = sanitized.strip('_')
    
    return sanitized


def ensure_template_exists(template_path: str = None) -> str:
    """
    Ensure the DOCX template exists, creating a minimal one if necessary.
    
    Args:
        template_path: Path to template, defaults to templates/word/ReportTemplate.docx
        
    Returns:
        Path to the template file
    """
    if template_path is None:
        # Use project root templates directory
        project_root = Path(__file__).parent.parent.parent
        template_dir = project_root / "templates" / "word"
        template_dir.mkdir(parents=True, exist_ok=True)
        template_path = str(template_dir / "ReportTemplate.docx")
    
    if os.path.exists(template_path):
        logger.info(f"Using existing template: {template_path}")
        return template_path
    
    # Create minimal template
    logger.info(f"Creating minimal template at: {template_path}")
    _create_minimal_template(template_path)
    return template_path


def ensure_assignment_template_exists() -> str:
    """
    Ensure the assignment batch template exists, creating one if necessary.
    
    Returns:
        Path to the assignment template file
    """
    project_root = Path(__file__).parent.parent.parent
    template_dir = project_root / "templates" / "word"
    template_dir.mkdir(parents=True, exist_ok=True)
    template_path = str(template_dir / "assignment_compiled.docx")
    
    if os.path.exists(template_path):
        logger.info(f"Using existing assignment template: {template_path}")
        return template_path
    
    # Create assignment batch template
    logger.info(f"Creating assignment batch template at: {template_path}")
    _create_assignment_template(template_path)
    return template_path


def _create_minimal_template(template_path: str):
    """Create a minimal DOCX template for reports aligned with teacher view"""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()

    # 1) 抬头信息（页首）
    title = doc.add_heading('批阅作业 - {{ assignmentTitle }}（{{ studentName }}）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Basic info in a clean table format
    doc.add_heading('基本信息', level=1)
    if DOCXTPL_AVAILABLE:
        # Create table with template variables
        doc.add_paragraph("""
{%- set info_data = [
    ('作业', assignmentTitle),
    ('学生', studentName), 
    ('提交时间', submittedAt)
] -%}
{% for label, value in info_data %}{{ label }}：{{ value }}
{% endfor %}
        """.strip())
    else:
        info_para = doc.add_paragraph()
        info_para.add_run('作业：').bold = True
        info_para.add_run('[作业标题]\n')
        info_para.add_run('学生：').bold = True
        info_para.add_run('[学生姓名]\n')
        info_para.add_run('提交时间：').bold = True
        info_para.add_run('[提交时间]')

    # 2) 总分与维度评分表
    doc.add_heading('评分结果', level=1)
    score_para = doc.add_paragraph()
    score_para.add_run('总分：').bold = True
    if DOCXTPL_AVAILABLE:
        score_para.add_run('{{ gradingResult.total_score|default(0) }} / {{ total_max_score|default(40) }}')
    else:
        score_para.add_run('[总分]')

    # 维度明细表 - Using proper table instead of markdown
    doc.add_heading('维度评分明细', level=2)
    if DOCXTPL_AVAILABLE:
        # Create a proper DOCX table using template
        doc.add_paragraph("""
{% if gradingResult.dimensions and gradingResult.dimensions|length > 0 %}
{%- set table_data = [] -%}
{%- for dim in gradingResult.dimensions -%}
{%- set _ = table_data.append([dim.dimension_name, dim.score, dim.selected_rubric_level|default(''), dim.feedback|default('')]) -%}
{%- endfor -%}
{% for row in table_data %}{{ row[0] }}	{{ row[1] }}	{{ row[2] }}	{{ row[3] }}
{% endfor %}

{% for dim in gradingResult.dimensions %}
{{ dim.dimension_name }}维度详情：

亮点句子：
{% if dim.example_good_sentence and dim.example_good_sentence|length > 0 %}
{% for sentence in dim.example_good_sentence %}
• {{ sentence }}
{% endfor %}
{% else %}
• 无
{% endif %}

待改进句：
{% if dim.example_improvement_suggestion and dim.example_improvement_suggestion|length > 0 %}
{% for suggestion in dim.example_improvement_suggestion %}
- 原文：{{ suggestion.original|default('') }}
- 建议：{{ suggestion.suggested|default('') }}
{% endfor %}
{% else %}
• 无
{% endif %}

{% endfor %}
{% else %}
本次作文评估采用系统性标准，重点关注内容理解、结构组织、语言表达和文采创新等维度。建议继续加强写作练习以提升各项能力。
{% endif %}
        """.strip())
    else:
        doc.add_paragraph('[维度评分明细]')

    # 3) 作文正文（当前文本）
    doc.add_heading('作文正文', level=1)
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph('{{ currentEssayContent|default("作文内容将在此处显示。建议学生认真审题，组织好文章结构，表达清楚完整的思想。") }}')
    else:
        doc.add_paragraph('[作文正文内容]')

    # 3.1) 作文图片（如果有的话）
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph('{% if images.primary_image_path %}')
        doc.add_heading('作文图片', level=2)
        # Use InlineImage objects directly if available
        doc.add_paragraph('{% if images.primary_image %}{{ images.primary_image }}{% elif images.friendly_message %}{{ images.friendly_message }}{% else %}图片缺失或不可访问{% endif %}')
        doc.add_paragraph('{% endif %}')
    else:
        doc.add_heading('作文图片', level=2)
        doc.add_paragraph('[作文图片]')

    # 4) 综合评价与寄语
    doc.add_heading('综合评价与寄语', level=1)
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph("""
{% if gradingResult.overall_comment %}
{{ gradingResult.overall_comment }}
{% else %}
{% if scores.total >= 32 %}
本次作文总体表现良好，获得{{ scores.total }}分，显示了扎实的写作基础和良好的表达能力。
{% elif scores.total >= 24 %}
本次作文表现中等，获得{{ scores.total }}分，在某些方面表现出色，但仍有进一步提升的空间。
{% elif scores.total > 0 %}
本次作文获得{{ scores.total }}分，需要在多个方面加强练习，建议重点关注写作基础技能的提升。
{% else %}
本次作文体现了一定的写作基础，建议继续加强练习，在结构组织和语言表达方面进一步提升。
{% endif %}
{% endif %}
        """.strip())
    else:
        doc.add_paragraph('[综合评价与寄语]')

    # 5) 主要优点
    doc.add_heading('主要优点', level=1)
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph("""
{% if gradingResult.strengths and gradingResult.strengths|length > 0 %}
{% for strength in gradingResult.strengths %}
• {{ strength }}
{% endfor %}
{% else %}
• 能够完成作文基本要求
• 语言表达基本流畅
• 内容结构相对完整
{% endif %}
        """.strip())
    else:
        doc.add_paragraph('[主要优点]')

    # 6) 改进建议
    doc.add_heading('改进建议', level=1)
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph("""
{% if gradingResult.improvements and gradingResult.improvements|length > 0 %}
{% for improvement in gradingResult.improvements %}
• {{ improvement }}
{% endfor %}
{% else %}
• 可以进一步丰富内容深度
• 语言表达可以更加精准
• 文章结构可以更加紧密
{% endif %}
        """.strip())
    else:
        doc.add_paragraph('[改进建议]')

    # 7) AI 增强内容审核
    doc.add_heading('AI 增强内容审核', level=1)

    # 段落大纲分析
    doc.add_heading('段落大纲分析', level=2)
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph("""
{% if outline and outline|length > 0 %}
{% for item in outline %}
{{ item.index }}. {{ item.intention }}
{% endfor %}
{% else %}
本次作文结构分析：建议注意段落之间的逻辑关系，确保文章结构清晰，层次分明。
{% endif %}
        """.strip())
    else:
        doc.add_paragraph('[段落大纲分析]')

    # 诊断建议
    doc.add_heading('诊断建议', level=2)
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph("""
{% if diagnoses and diagnoses|length > 0 %}
{% for diag in diagnoses %}
{{ diag.id }}. {{ diag.target }} - {{ diag.evidence }} 
   建议：{% for suggestion in diag.suggestions %}{{ suggestion }}{% if not loop.last %}；{% endif %}{% endfor %}
{% endfor %}
{% else %}
建议重点关注：1. 加强审题能力；2. 提升语言表达的准确性；3. 增强文章的逻辑性和条理性。
{% endif %}
        """.strip())
    else:
        doc.add_paragraph('[诊断建议]')

    # 个性化练习
    doc.add_heading('个性化练习', level=2)
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph("""
{% if personalizedPractices and personalizedPractices|length > 0 %}
{% for practice in personalizedPractices %}
{{ loop.index }}. {{ practice.title }}
   要求：{{ practice.requirement }}
{% endfor %}
{% else %}
推荐练习：1. 每日阅读优秀文章并摘录好词好句；2. 练习写作文提纲；3. 加强审题训练，确保文章切题。
{% endif %}
        """.strip())
    else:
        doc.add_paragraph('[个性化练习]')

    # 综合诊断总结
    doc.add_heading('综合诊断总结', level=2)
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph("""
{% if summaryData %}
问题总结：{{ summaryData.problemSummary|default("本次作文分析发现的主要问题包括结构组织、语言表达等方面。") }}

改进建议：{{ summaryData.improvementPlan|default("建议从基础写作技巧、段落结构、词汇运用等方面进行针对性改进。") }}

预期效果：{{ summaryData.expectedOutcome|default("通过有针对性的练习和指导，预期能够在作文质量上取得明显提升。") }}
{% else %}
问题总结：本次作文分析发现的主要问题包括结构组织、语言表达等方面。

改进建议：建议从基础写作技巧、段落结构、词汇运用等方面进行针对性改进。

预期效果：通过有针对性的练习和指导，预期能够在作文质量上取得明显提升。
{% endif %}
        """.strip())
    else:
        doc.add_paragraph('[综合诊断总结]')

    # 给家长的总结
    doc.add_heading('给家长的总结', level=2)
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph('{% if parentSummary %}{{ parentSummary }}{% else %}总体而言，该作文具有一定的优点，同时也存在一些需要改进的地方。建议家长鼓励孩子多读多写，持续提升写作能力。{% endif %}')
    else:
        doc.add_paragraph('[给家长的总结]')

    # Footer
    doc.add_paragraph().add_run('\n' + '='*50)
    footer = doc.add_paragraph()
    footer.add_run('报告生成时间: ').bold = True
    if DOCXTPL_AVAILABLE:
        footer.add_run('{{ now|strftime("%Y-%m-%d %H:%M:%S") }}')
    else:
        footer.add_run('[生成时间]')
    footer.alignment = 1  # Center

    doc.save(template_path)
    logger.info(f"Created teacher-view aligned template: {template_path}")
def _create_assignment_template(template_path: str):
    """Create assignment batch template with student loop and page breaks"""
    from docx import Document

    doc = Document()

    # Assignment header
    title = doc.add_heading('{{ assignment.title }} 批量作文评估报告', 0)
    title.alignment = 1  # Center

    # Assignment metadata
    doc.add_heading('作业信息', level=1)
    info_para = doc.add_paragraph()
    info_para.add_run('作业标题：').bold = True
    info_para.add_run('{{ assignment.title }}\n')
    info_para.add_run('班级：').bold = True
    # 仅显示名称，若已是字符串则直接显示
    info_para.add_run('{{ assignment.classroom.name|default(assignment.classroom) }}\n')
    info_para.add_run('教师：').bold = True
    info_para.add_run('{{ assignment.teacher.name|default(assignment.teacher) }}\n')
    info_para.add_run('生成时间：').bold = True
    if DOCXTPL_AVAILABLE:
        info_para.add_run('{{ now|strftime("%Y年%m月%d日 %H:%M:%S") }}')
    else:
        info_para.add_run('{{ current_time }}')

    # Add a page break before student reports
    doc.add_page_break()

    # Student reports loop
    if DOCXTPL_AVAILABLE:
        doc.add_paragraph('{% for s in students %}')

        # Student page header
        doc.add_heading('{{ s.student_name }} 作文评估报告', level=1)

        # Basic student info
        basic_info = doc.add_paragraph()
        basic_info.add_run('学生：').bold = True
        basic_info.add_run('{{ s.student_name }}\n')
        basic_info.add_run('题目：').bold = True
        basic_info.add_run('{{ s.topic }}\n')
        basic_info.add_run('字数：').bold = True
        basic_info.add_run('{{ s.words|default("未统计") }}\n')
        basic_info.add_run('总分：').bold = True
        basic_info.add_run('{{ s.scores.total }}')

        # Scoring dimensions table — 使用中括号访问 key，避免与 dict.items 冲突
        doc.add_heading('评分维度', level=2)
        doc.add_paragraph("""
{% if s.scores["items"] and s.scores["items"]|length > 0 %}
| 维度 | 得分 | 满分 | 权重 | 理由 |
|------|------|------|------|------|
{% for i in s.scores["items"] %}| {{ i.name }} | {{ i.score }} | {{ i.max_score }} | {{ i.weight|default('') }} | {{ i.reason|default('') }} |
{% endfor %}{% else %}（暂无评分维度数据）{% endif %}
        """.strip())

        # Cleaned text
        doc.add_heading('清洗后文本', level=2)
        doc.add_paragraph('{{ s.text.cleaned|default("") }}')

        # Analysis and diagnostics
        doc.add_heading('分析与诊断', level=2)

        # Structure analysis
        doc.add_heading('结构分析', level=3)
        doc.add_paragraph("""
{% for o in s.analysis.outline %}
第{{ o.para }}段：{{ o.intent }}
{% endfor %}
        """.strip())

        # Issue list
        doc.add_heading('问题清单', level=3)
        doc.add_paragraph("""
{% for iss in s.analysis.issues %}
- {{ iss }}
{% endfor %}
        """.strip())

        # Diagnostic suggestions — 用 if/else 包裹，避免 for-else 歧义
        doc.add_heading('诊断建议', level=3)
        doc.add_paragraph("""
{% if s.diagnostics and s.diagnostics|length > 0 %}
{% for d in s.diagnostics %}
{% if d.para %}第{{ d.para }}段{% else %}全文{% endif %}｜{{ d.issue }}｜证据：{{ d.evidence }}｜建议：{{ d.advice|default([])|join('；') }}
{% endfor %}
{% else %}（暂无诊断建议）{% endif %}
        """.strip())

        # Diagnosis summary
        doc.add_heading('诊断总结', level=3)
        doc.add_paragraph('{{ s.diagnosis.before|default("") }}')
        doc.add_paragraph('{{ s.diagnosis.comment|default("") }}')
        doc.add_paragraph('{{ s.diagnosis.after|default("") }}')
        doc.add_paragraph('{{ s.summary|default("") }}')

        # Paragraph-level enhancements
        doc.add_heading('段落级增强', level=2)
        doc.add_paragraph("""
{% for p in s.paragraphs %}
第{{ p.para_num }}段：意图：{{ p.intent|default('') }}
原文：{{ p.original_text|default('') }}
反馈：{{ p.feedback|default('') }}
优化：{{ p.polished_text|default('') }}

{% endfor %}
        """.strip())

        # Personalized exercises — 兼容 hint/hints
        doc.add_heading('个性化练习', level=2)
        doc.add_paragraph("""
{% if s.exercises and s.exercises|length > 0 %}
{% for ex in s.exercises %}
[{{ ex.type }}] {{ ex.prompt }}
要点：{{ (ex.hint if ex.hint is not none else ex.hints)|default([])|join('；') }}
示例：{{ ex.sample|default('') }}

{% endfor %}{% else %}（暂无个性化练习）{% endif %}
        """.strip())

        # Images section — 修正图片显示逻辑
        doc.add_paragraph('{% if s.images.primary_image_path %}')
        doc.add_heading('作文图片', level=2)
        # Use InlineImage objects directly if available
        doc.add_paragraph('{% if s.images.primary_image %}{{ s.images.primary_image }}{% elif s.images.friendly_message %}{{ s.images.friendly_message }}{% else %}图片缺失或不可访问{% endif %}')
        doc.add_paragraph('{% endif %}')
        doc.add_paragraph('{% endif %}')
        doc.add_paragraph('{% endif %}')

        # Page break between students (except for the last one)
        doc.add_paragraph('{% if not loop.last %}')
        doc.add_page_break()
        doc.add_paragraph('{% endif %}')

        # End student loop
        doc.add_paragraph('{% endfor %}')
    else:
        # Fallback for non-docxtpl
        doc.add_paragraph('[学生报告循环 - 需要 docxtpl 支持]')

    # Footer
    doc.add_paragraph().add_run('\n' + '='*50)
    footer = doc.add_paragraph()
    footer.add_run('报告生成：e文智教系统 - ').bold = True
    if DOCXTPL_AVAILABLE:
        footer.add_run('{{ now|strftime("%Y年%m月%d日") }}')
    else:
        footer.add_run('{{ current_time }}')
    footer.alignment = 1  # Center

    doc.save(template_path)
def render_essay_docx(evaluation: EvaluationResult, output_path: str = None, review_status: str = None, teacher_view: bool = False) -> str:
    """
    Render a single essay evaluation to DOCX.
    
    Args:
        evaluation: EvaluationResult instance
        output_path: Output file path, auto-generated if None
        review_status: Review status for display (ai_generated, teacher_reviewed, finalized)
        teacher_view: If True, use teacher view aligned template structure
        
    Returns:
        Path to generated DOCX file
    """
    if output_path is None:
        # Generate filename from student name and topic
        student = _sanitize_filename(evaluation.studentName or evaluation.meta.student)
        topic = _sanitize_filename(evaluation.assignmentTitle or str(evaluation.meta.topic))
        date_str = _sanitize_filename(str(evaluation.meta.date))
        filename = f"{student}_{topic}_{date_str}.docx"
        
        # Use temp directory
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, filename)
    
    # Auto-detect teacher view mode if new fields are present
    if not teacher_view:
        teacher_view = (evaluation.assignmentTitle is not None or 
                       evaluation.currentEssayContent is not None or
                       evaluation.outline or evaluation.diagnoses)
    
    if DOCXTPL_AVAILABLE:
        return _render_with_docxtpl(evaluation, output_path, review_status, teacher_view)
    else:
        return _render_with_python_docx(evaluation, output_path, review_status, teacher_view)


def render_assignment_docx(assignment_id: int, evaluations: list = None, output_path: str = None) -> str:
    """
    Render assignment summary DOCX.
    
    Args:
        assignment_id: Assignment ID
        evaluations: List of EvaluationResult instances
        output_path: Output file path
        
    Returns:
        Path to generated DOCX file
        
    Raises:
        FileNotFoundError: If assignment template is missing
        ValueError: If no evaluations available
    """
    # Check if assignment template exists
    project_root = Path(__file__).parent.parent.parent
    assignment_template_path = project_root / "templates" / "word" / "assignment_compiled.docx"
    
    if not assignment_template_path.exists():
        raise FileNotFoundError(
            f"Assignment template not found at {assignment_template_path}. "
            "Assignment summary export requires a dedicated template. "
            "Please create the template or use individual essay export instead."
        )
    
    if not evaluations or len(evaluations) == 0:
        raise ValueError(f"No evaluation data available for assignment {assignment_id}")
    
    logger.info(f"Rendering assignment summary for {len(evaluations)} evaluations")
    
    # For now, use the existing combined rendering logic if available
    # This is a placeholder for proper assignment template implementation
    try:
        return _render_assignment_combined(assignment_id, evaluations, output_path)
    except NotImplementedError:
        # If combined rendering is not implemented, provide clear guidance
        raise NotImplementedError(
            f"Assignment summary rendering not yet fully implemented. "
            f"Available options: 1) Create individual essay reports, "
            f"2) Implement _render_assignment_combined function, "
            f"3) Add assignment template at {assignment_template_path}"
        )


def _render_assignment_combined(assignment_id: int, evaluations: list, output_path: str = None) -> str:
    """
    Render combined assignment DOCX using assignment template.
    
    This is a placeholder implementation - should be enhanced based on actual requirements.
    """
    raise NotImplementedError("Combined assignment rendering not yet implemented")


def _render_with_docxtpl(evaluation: EvaluationResult, output_path: str, review_status: str = None, teacher_view: bool = False) -> str:
    """Render using docxtpl (template-based)"""
    template_path = ensure_template_exists()
    
    try:
        doc = DocxTemplate(template_path)
        context = to_context(evaluation, doc_template=doc)
        
        # Add current timestamp and enhance context
        from datetime import datetime
        context['now'] = datetime.now()
        context['current_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Add review status information
        context['review_status'] = review_status or 'ai_generated'
        context['is_reviewed'] = review_status in ['teacher_reviewed', 'finalized'] if review_status else False
        context['needs_review_warning'] = review_status not in ['teacher_reviewed', 'finalized'] if review_status else True
        
        # Add fields for future enhancements (P2)
        context.setdefault('paragraphs', [])
        context.setdefault('exercises', [])
        context.setdefault('feedback_summary', '')
        
        # For teacher view mode, ensure all required fields are populated with defaults
        if teacher_view:
            context.setdefault('total_max_score', 40)  # Default max score
            
            # Ensure gradingResult is properly structured
            if 'gradingResult' not in context:
                context['gradingResult'] = {
                    'total_score': context.get('scores', {}).get('total', 0),
                    'dimensions': [],
                    'overall_comment': context.get('overall_comment', ''),
                    'strengths': context.get('strengths', []),
                    'improvements': context.get('improvements', [])
                }
        
        # Create Jinja environment with custom filters (P0)
        from jinja2 import Environment
        env = Environment(autoescape=False)
        
        def strftime_filter(dt, fmt):
            """Custom strftime filter that handles both datetime objects and strings"""
            if dt is None:
                return ''
            if isinstance(dt, str):
                return dt  # If already a string, return as-is
            if hasattr(dt, 'strftime'):
                return dt.strftime(fmt)
            return str(dt)
        
        env.filters['strftime'] = strftime_filter
        
        # Render with custom jinja environment
        doc.render(context, jinja_env=env)
        doc.save(output_path)
        
        logger.info(f"Rendered DOCX using docxtpl: {output_path}")
        return output_path
        
    except FileNotFoundError as e:
        logger.info(f"Template file not found, falling back to python-docx: {e}")
        return _render_with_python_docx(evaluation, output_path, review_status, teacher_view)
    except Exception as e:
        # P1: Don't fallback on template syntax errors - raise them clearly
        logger.error(f"Failed to render with docxtpl due to template error: {e}")
        raise RuntimeError(f"DOCX template rendering failed: {e}") from e


def _render_with_python_docx(evaluation: EvaluationResult, output_path: str, review_status: str = None, teacher_view: bool = False) -> str:
    """Render using python-docx (direct generation) with both legacy and teacher view support"""
    doc = Document()
    
    # Check if we should use teacher view structure
    use_teacher_view = teacher_view or (evaluation.assignmentTitle is not None or 
                                       evaluation.currentEssayContent is not None)
    
    if use_teacher_view:
        # Teacher view aligned structure
        _render_teacher_view_structure(doc, evaluation, review_status)
    else:
        # Legacy structure
        _render_legacy_structure(doc, evaluation, review_status)
    
    doc.save(output_path)
    logger.info(f"Rendered DOCX using python-docx: {output_path}")
    return output_path


def _render_teacher_view_structure(doc, evaluation: EvaluationResult, review_status: str = None):
    """Render teacher view aligned structure using python-docx"""
    # 1) 抬头信息（页首）
    assignment_title = evaluation.assignmentTitle or evaluation.meta.topic or "未知作业"
    student_name = evaluation.studentName or evaluation.meta.student or "未知学生"
    title = doc.add_heading(f'批阅作业 - {assignment_title}（{student_name}）', 0)
    title.alignment = 1  # Center
    
    # Add review status warning if needed
    if review_status and review_status != 'teacher_reviewed' and review_status != 'finalized':
        warning = doc.add_paragraph()
        warning.add_run('⚠️ 注意：此报告内容为AI生成，尚未经过教师审核确认').bold = True
        warning.style = 'Intense Quote'
    
    # Basic information
    doc.add_heading('基本信息', level=1)
    basic_info = doc.add_paragraph()
    basic_info.add_run('作业：').bold = True
    basic_info.add_run(f'{assignment_title}\n')
    basic_info.add_run('学生：').bold = True
    basic_info.add_run(f'{student_name}\n')
    basic_info.add_run('提交时间：').bold = True
    submitted_at = evaluation.submittedAt or str(evaluation.meta.date) or "未知时间"
    basic_info.add_run(submitted_at)
    
    # 2) 总分与维度评分表
    doc.add_heading('评分结果', level=1)
    scores_para = doc.add_paragraph()
    scores_para.add_run('总分：').bold = True
    total_score = evaluation.scores.total if evaluation.scores else 0
    scores_para.add_run(f'{total_score}分\n')
    
    # 维度明细表
    doc.add_heading('维度评分明细', level=2)
    if evaluation.scores and evaluation.scores.rubrics:
        # Create proper DOCX table
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        table = doc.add_table(rows=1 + len(evaluation.scores.rubrics), cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Set table headers
        headers = ['维度', '分数', '等级', '维度反馈']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fill table data
        for row_idx, rubric in enumerate(evaluation.scores.rubrics, 1):
            table.cell(row_idx, 0).text = rubric.name
            table.cell(row_idx, 1).text = str(rubric.score)
            table.cell(row_idx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row_idx, 2).text = getattr(rubric, 'level', '')
            table.cell(row_idx, 3).text = getattr(rubric, 'reason', '')
            
            # Add dimension details section after table
        
        # Add dimension details sections
        for rubric in evaluation.scores.rubrics:
            dim_detail = doc.add_heading(f'{rubric.name}维度详情：', level=3)
            
            # Bright points (placeholder - would need to be populated from actual data)
            bright_para = doc.add_paragraph()
            bright_para.add_run('亮点句子：').bold = True
            
            # Check if this is an AI-enhanced evaluation (has AI-generated enhancement fields)
            is_ai_enhanced = (hasattr(evaluation, 'strengths') and evaluation.strengths) or \
                            (hasattr(evaluation, 'improvements') and evaluation.improvements) or \
                            (hasattr(evaluation, 'overall_comment') and evaluation.overall_comment)
            
            # Check if rubric has example_good_sentence data
            if hasattr(rubric, 'example_good_sentence') and rubric.example_good_sentence:
                for sentence in rubric.example_good_sentence:
                    doc.add_paragraph(f'• {sentence}')
            elif not is_ai_enhanced:
                # Only show fallback if this is not an AI-enhanced evaluation
                doc.add_paragraph('• 无')
            # If AI-enhanced but no example sentences, don't show anything
            
            # Improvement suggestions
            improve_para = doc.add_paragraph()
            improve_para.add_run('待改进句：').bold = True
            
            # Check if rubric has example_improvement_suggestion data
            if hasattr(rubric, 'example_improvement_suggestion') and rubric.example_improvement_suggestion:
                for suggestion in rubric.example_improvement_suggestion:
                    original = getattr(suggestion, 'original', '') if hasattr(suggestion, 'original') else suggestion.get('original', '') if isinstance(suggestion, dict) else ''
                    suggested = getattr(suggestion, 'suggested', '') if hasattr(suggestion, 'suggested') else suggestion.get('suggested', '') if isinstance(suggestion, dict) else ''
                    if original and suggested:
                        doc.add_paragraph(f'- 原文：{original}\n- 建议：{suggested}')
                    else:
                        doc.add_paragraph(f'• {suggestion}')
            elif not is_ai_enhanced:
                # Only show fallback if this is not an AI-enhanced evaluation
                doc.add_paragraph('• 无')
            # If AI-enhanced but no improvement suggestions, don't show anything
    else:
        # Provide meaningful fallback for rubrics when no scoring data
        doc.add_paragraph('本次作文评估采用系统性标准，重点关注内容理解、结构组织、语言表达和文采创新等维度。建议继续加强写作练习以提升各项能力。')
    
    # 3) 作文正文（当前文本）
    doc.add_heading('作文正文', level=1)
    essay_content = evaluation.currentEssayContent
    if not essay_content and evaluation.text:
        essay_content = evaluation.text.cleaned or evaluation.text.original
    if not essay_content:
        essay_content = "作文内容将在此处显示。建议学生认真审题，组织好文章结构，表达清楚完整的思想。"
    doc.add_paragraph(essay_content)
    
    # 3.1) 作文图片（如果有的话）
    if hasattr(evaluation, '_essay_instance'):
        essay = evaluation._essay_instance
        if essay and essay.original_image_path:
            doc.add_heading('作文图片', level=2)
            
            # Try to add the composed image (original + annotations)
            try:
                from pathlib import Path
                from docx.shared import Inches
                from app.reporting.image_overlay import compose_overlay_images
                from app.utils.path_resolver import resolve_upload_path, get_friendly_image_message
                
                # Resolve the original image path
                resolved_original_path = resolve_upload_path(essay.original_image_path)
                
                if resolved_original_path:
                    composed_image_path = None
                    
                    # If we have both original and overlay, compose them
                    if essay.annotated_overlay_path:
                        resolved_overlay_path = resolve_upload_path(essay.annotated_overlay_path)
                        if resolved_overlay_path:
                            composed_image_path = compose_overlay_images(resolved_original_path, resolved_overlay_path)
                            if composed_image_path:
                                doc.add_paragraph('作文图片（含教师批注）：')
                                doc.add_picture(composed_image_path, width=Inches(6))
                            else:
                                # Fallback: show original and overlay separately if composition failed
                                doc.add_paragraph('原始图片：')
                                doc.add_picture(resolved_original_path, width=Inches(6))
                                doc.add_paragraph('教师批注图片：')
                                doc.add_picture(resolved_overlay_path, width=Inches(6))
                        else:
                            # Only original image exists, overlay path can't be resolved
                            doc.add_paragraph('作文图片：')
                            doc.add_picture(resolved_original_path, width=Inches(6))
                    else:
                        # Only original image exists, no annotations
                        doc.add_paragraph('作文图片：')
                        doc.add_picture(resolved_original_path, width=Inches(6))
                        
                else:
                    doc.add_paragraph(get_friendly_image_message())
            except Exception as e:
                logger.warning(f"Failed to add image to DOCX: {e}")
                doc.add_paragraph('图片加载失败')
    
    # 4) 综合评价与寄语
    doc.add_heading('综合评价与寄语', level=1)
    overall_comment = getattr(evaluation, 'overall_comment', '') or ""
    if not overall_comment:
        # Provide meaningful fallback based on score if available
        total_score = getattr(evaluation.scores, 'total', 0) if hasattr(evaluation, 'scores') and evaluation.scores else 0
        if total_score >= 32:  # Assuming 40 is max, 80% threshold
            overall_comment = f"本次作文总体表现良好，获得{total_score}分，显示了扎实的写作基础和良好的表达能力。"
        elif total_score >= 24:  # 60% threshold
            overall_comment = f"本次作文表现中等，获得{total_score}分，在某些方面表现出色，但仍有进一步提升的空间。"
        elif total_score > 0:
            overall_comment = f"本次作文获得{total_score}分，需要在多个方面加强练习，建议重点关注写作基础技能的提升。"
        else:
            overall_comment = "本次作文体现了一定的写作基础，建议继续加强练习，在结构组织和语言表达方面进一步提升。"
    doc.add_paragraph(overall_comment)
    
    # 5) 主要优点
    doc.add_heading('主要优点', level=1)
    strengths = getattr(evaluation, 'strengths', [])
    if not strengths:
        # Provide meaningful fallback strengths
        strengths = ["能够完成作文基本要求", "语言表达基本流畅", "内容结构相对完整"]
    for strength in strengths:
        strength_para = doc.add_paragraph()
        strength_para.add_run(f'• {strength}')
    
    # 6) 改进建议
    doc.add_heading('改进建议', level=1)
    improvements = getattr(evaluation, 'improvements', [])
    if not improvements:
        # Provide meaningful fallback improvements
        improvements = ["可以进一步丰富内容深度", "语言表达可以更加精准", "文章结构可以更加紧密"]
    for improvement in improvements:
        improve_para = doc.add_paragraph()
        improve_para.add_run(f'• {improvement}')
    
    # 7) AI 增强内容审核
    doc.add_heading('AI 增强内容审核', level=1)
    
    # 段落大纲分析
    doc.add_heading('段落大纲分析', level=2)
    if evaluation.outline:
        for item in evaluation.outline:
            outline_para = doc.add_paragraph()
            outline_para.add_run(f'{item.get("index", 0)}. {item.get("intention", "")}')
    else:
        # Provide meaningful fallback for outline
        doc.add_paragraph('本次作文结构分析：建议注意段落之间的逻辑关系，确保文章结构清晰，层次分明。')
    
    # 诊断建议
    doc.add_heading('诊断建议', level=2)
    if evaluation.diagnoses:
        for diag in evaluation.diagnoses:
            diag_para = doc.add_paragraph()
            diag_id = diag.get('id', 0)
            target = diag.get('target', '')
            evidence = diag.get('evidence', '')
            suggestions = '；'.join(diag.get('suggestions', []))
            diag_para.add_run(f'{diag_id}. {target} - {evidence}')
            if suggestions:
                diag_para.add_run(f'\n   建议：{suggestions}')
    else:
        # Provide meaningful fallback for diagnoses
        doc.add_paragraph('建议重点关注：1. 加强审题能力；2. 提升语言表达的准确性；3. 增强文章的逻辑性和条理性。')
    
    # 个性化练习
    doc.add_heading('个性化练习', level=2)
    if evaluation.personalizedPractices:
        for i, practice in enumerate(evaluation.personalizedPractices, 1):
            practice_para = doc.add_paragraph()
            title = practice.get('title', '')
            requirement = practice.get('requirement', '')
            practice_para.add_run(f'{i}. {title}')
            if requirement:
                practice_para.add_run(f'\n   要求：{requirement}')
    else:
        # Provide meaningful fallback for practices
        doc.add_paragraph('推荐练习：1. 每日阅读优秀文章并摘录好词好句；2. 练习写作文提纲；3. 加强审题训练，确保文章切题。')
    
    # 综合诊断总结
    doc.add_heading('综合诊断总结', level=2)
    if evaluation.summaryData:
        summary_para = doc.add_paragraph()
        summary_para.add_run('问题总结：').bold = True
        problem_summary = evaluation.summaryData.get("problemSummary", "") or "本次作文分析发现的主要问题包括结构组织、语言表达等方面。"
        summary_para.add_run(f'{problem_summary}\n\n')
        summary_para.add_run('改进建议：').bold = True
        improvement_plan = evaluation.summaryData.get("improvementPlan", "") or "建议从基础写作技巧、段落结构、词汇运用等方面进行针对性改进。"
        summary_para.add_run(f'{improvement_plan}\n\n')
        summary_para.add_run('预期效果：').bold = True
        expected_outcome = evaluation.summaryData.get("expectedOutcome", "") or "通过有针对性的练习和指导，预期能够在作文质量上取得明显提升。"
        summary_para.add_run(f'{expected_outcome}')
    else:
        # Provide meaningful fallback for summary data
        summary_para = doc.add_paragraph()
        summary_para.add_run('问题总结：').bold = True
        summary_para.add_run('本次作文分析发现的主要问题包括结构组织、语言表达等方面。\n\n')
        summary_para.add_run('改进建议：').bold = True
        summary_para.add_run('建议从基础写作技巧、段落结构、词汇运用等方面进行针对性改进。\n\n')
        summary_para.add_run('预期效果：').bold = True
        summary_para.add_run('通过有针对性的练习和指导，预期能够在作文质量上取得明显提升。')
    
    # 给家长的总结
    doc.add_heading('给家长的总结', level=2)
    parent_summary = evaluation.parentSummary or ""
    if not parent_summary:
        # Provide meaningful fallback for parent summary
        parent_summary = "总体而言，该作文具有一定的优点，同时也存在一些需要改进的地方。建议家长鼓励孩子多读多写，持续提升写作能力。"
    doc.add_paragraph(parent_summary)
    
    # Footer
    doc.add_paragraph().add_run('\n' + '='*50)
    footer = doc.add_paragraph()
    footer.add_run('报告生成时间：').bold = True
    from datetime import datetime
    footer.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    footer.alignment = 1  # Center


def _render_legacy_structure(doc, evaluation: EvaluationResult, review_status: str = None):
    """Render legacy structure using python-docx"""
    # Title
    title = doc.add_heading(f'{evaluation.meta.topic} 作文评估报告', 0)
    title.alignment = 1  # Center
    
    # Add review status warning if needed
    if review_status and review_status != 'teacher_reviewed' and review_status != 'finalized':
        warning = doc.add_paragraph()
        warning.add_run('⚠️ 注意：此报告内容为AI生成，尚未经过教师审核确认').bold = True
        warning.style = 'Intense Quote'
    
    # Basic information
    doc.add_heading('基本信息', level=1)
    basic_info = doc.add_paragraph()
    basic_info.add_run('学生：').bold = True
    basic_info.add_run(f'{evaluation.meta.student}\n')
    basic_info.add_run('班级：').bold = True
    basic_info.add_run(f'{evaluation.meta.class_}\n')
    basic_info.add_run('教师：').bold = True
    basic_info.add_run(f'{evaluation.meta.teacher}\n')
    basic_info.add_run('日期：').bold = True
    basic_info.add_run(f'{evaluation.meta.date}')
    
    # Scores
    doc.add_heading('评分结果', level=1)
    scores_para = doc.add_paragraph()
    scores_para.add_run('总分：').bold = True
    scores_para.add_run(f'{evaluation.scores.total}分\n')
    
    # Rubrics detail
    if evaluation.scores.rubrics:
        doc.add_heading('各维度评分', level=2)
        for rubric in evaluation.scores.rubrics:
            rubric_para = doc.add_paragraph()
            rubric_para.add_run(f'{rubric.name}：').bold = True
            rubric_para.add_run(f'{rubric.score}/{rubric.max}分 (权重{rubric.weight})\n')
            if rubric.reason:
                rubric_para.add_run(f'理由：{rubric.reason}')
    
    # Original text
    if evaluation.text:
        doc.add_heading('原文', level=1)
        doc.add_paragraph(evaluation.text.original or '无原文内容')
        
        # Cleaned text
        if evaluation.text.cleaned and evaluation.text.cleaned != evaluation.text.original:
            doc.add_heading('清洗后文本', level=1)
            doc.add_paragraph(evaluation.text.cleaned)
    
    # Highlights summary
    if evaluation.highlights:
        doc.add_heading('高亮摘要', level=1)
        from app.schemas.evaluation import to_context
        context = to_context(evaluation)
        if 'highlight_summary' in context:
            for highlight_type, data in context['highlight_summary'].items():
                highlight_para = doc.add_paragraph()
                highlight_para.add_run(f'{highlight_type}：').bold = True
                highlight_para.add_run(f'低{data["low"]}个, 中{data["medium"]}个, 高{data["high"]}个\n')
                if data.get('examples'):
                    highlight_para.add_run('示例：').bold = True
                    highlight_para.add_run(', '.join(data['examples']))
    
    # Diagnosis
    if evaluation.diagnosis:
        doc.add_heading('诊断建议', level=1)
        if evaluation.diagnosis.before:
            doc.add_paragraph(evaluation.diagnosis.before)
        if evaluation.diagnosis.comment:
            doc.add_paragraph(evaluation.diagnosis.comment)
        if evaluation.diagnosis.after:
            doc.add_paragraph(evaluation.diagnosis.after)
    
    # Legacy diagnostics for backward compatibility
    if evaluation.diagnostics:
        doc.add_heading('诊断详情', level=1)
        for diag in evaluation.diagnostics:
            diag_para = doc.add_paragraph()
            para_info = f"第{diag.para}段" if diag.para else "全文"
            diag_para.add_run(f'{para_info} - {diag.issue}：').bold = True
            diag_para.add_run(f'{diag.evidence}\n')
            if diag.advice:
                diag_para.add_run('建议：').bold = True
                diag_para.add_run(', '.join(diag.advice))
    
    # Footer
    doc.add_paragraph().add_run('\n' + '='*50)
    footer = doc.add_paragraph()
    footer.add_run('报告生成时间：').bold = True
    from datetime import datetime
    footer.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    footer.alignment = 1  # Center