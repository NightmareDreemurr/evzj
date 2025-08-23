"""
Service layer for batch DOCX reporting.

This module provides business logic for building ViewModels and rendering
batch DOCX reports for assignments.
"""
import io
import tempfile
import logging
from typing import List, Optional, Literal, Union
from pathlib import Path

import zipstream
from app.extensions import db
from app.models import (
    EssayAssignment, Classroom, TeacherProfile, StudentProfile, 
    Essay, Enrollment
)
from app.dao.evaluation_dao import load_evaluation_by_essay
from app.services.evaluation_builder import load_evaluation_from_essay
from app.reporting.viewmodels import (
    StudentReportVM, AssignmentReportVM, ScoreVM,
    map_scores_to_vm, safe_get_student_name, safe_get_topic,
    safe_get_feedback, safe_get_original_paragraphs,
    map_paragraphs_to_vm, map_exercises_to_vm, build_feedback_summary
)
from app.reporting.docx_renderer import render_essay_docx
from app.schemas.evaluation import EvaluationResult, Meta, TextBlock, Scores, RubricScore

logger = logging.getLogger(__name__)


def build_student_vm(essay_id: int, require_review: bool = None) -> Optional[StudentReportVM]:
    """
    Build StudentReportVM from essay evaluation data.
    
    Args:
        essay_id: Essay ID
        require_review: Whether to require teacher review before export. 
                       If None, uses config EVAL_REQUIRE_REVIEW_BEFORE_EXPORT
        
    Returns:
        StudentReportVM instance or None if not found
        
    Raises:
        ValueError: If require_review is True but evaluation is not teacher-reviewed
    """
    try:
        from flask import current_app
        
        # Check if review is required
        if require_review is None:
            require_review = current_app.config.get('EVAL_REQUIRE_REVIEW_BEFORE_EXPORT', False)
        
        # Load evaluation result
        evaluation = load_evaluation_by_essay(essay_id)
        if not evaluation:
            logger.warning(f"No evaluation found for essay {essay_id}")
            return None
        
        # Get essay for additional context and status checking
        essay = db.session.get(Essay, essay_id)
        if not essay:
            logger.warning(f"Essay {essay_id} not found")
            return None
        
        # Check review status if required
        if require_review:
            eval_status = getattr(essay, 'evaluation_status', 'ai_generated')
            if eval_status not in ['teacher_reviewed', 'finalized']:
                error_msg = f"Essay {essay_id} evaluation status is '{eval_status}', but teacher review is required for export"
                logger.error(error_msg)
                raise ValueError(error_msg)
        
        # Add review status to meta if available
        eval_status = getattr(essay, 'evaluation_status', 'ai_generated')
        review_info = {
            'status': eval_status,
            'reviewed_by': getattr(essay, 'reviewed_by', None),
            'reviewed_at': getattr(essay, 'reviewed_at', None)
        }
        if not essay:
            logger.warning(f"Essay {essay_id} not found")
            return None
        
        # Extract student information
        student_id = 0
        student_name = safe_get_student_name(evaluation)
        student_no = None
        
        if essay.enrollment and essay.enrollment.student:
            student_profile = essay.enrollment.student
            student_id = student_profile.id
            if student_profile.user:
                student_name = student_profile.user.full_name or student_profile.user.username
            # Get student number from enrollment
            student_no = essay.enrollment.student_number
        
        # Extract other information
        topic = safe_get_topic(evaluation)
        words = getattr(evaluation.meta, 'words', None) if hasattr(evaluation, 'meta') else None
        scores = map_scores_to_vm(evaluation)
        feedback = safe_get_feedback(evaluation)
        original_paragraphs = safe_get_original_paragraphs(evaluation)
        
        # Extract new enhanced information
        paragraphs = map_paragraphs_to_vm(evaluation)
        exercises = map_exercises_to_vm(evaluation)
        feedback_summary = build_feedback_summary(evaluation)
        # For now, scanned_images is empty - can be populated from file storage later
        scanned_images = []
        
        return StudentReportVM(
            student_id=student_id,
            student_name=student_name,
            student_no=student_no,
            essay_id=essay_id,
            topic=topic,
            words=words,
            scores=scores,
            feedback=feedback,
            original_paragraphs=original_paragraphs,
            paragraphs=paragraphs,
            exercises=exercises,
            scanned_images=scanned_images,
            feedback_summary=feedback_summary,
            review_status=review_info['status'],
            reviewed_by=review_info['reviewed_by'],
            reviewed_at=review_info['reviewed_at'].isoformat() if review_info['reviewed_at'] else None
        )
        
    except Exception as e:
        logger.error(f"Failed to build student VM for essay {essay_id}: {e}")
        return None


def build_assignment_vm(assignment_id: int, require_review: bool = None) -> Optional[AssignmentReportVM]:
    """
    Build AssignmentReportVM with all student data.
    
    Args:
        assignment_id: Assignment ID
        require_review: Whether to require teacher review before export.
                       If None, uses config EVAL_REQUIRE_REVIEW_BEFORE_EXPORT
        
    Returns:
        AssignmentReportVM instance or None if not found
        
    Raises:
        ValueError: If require_review is True but some evaluations are not teacher-reviewed
    """
    try:
        # Get assignment with related data in single query to avoid N+1
        assignment = db.session.query(EssayAssignment)\
            .filter(EssayAssignment.id == assignment_id)\
            .first()
        
        if not assignment:
            logger.warning(f"Assignment {assignment_id} not found")
            return None
        
        # Get classroom and teacher info
        classroom_info = {"name": "未知班级", "id": None}
        teacher_info = {"name": "未知教师", "id": None}
        
        if assignment.teacher:
            teacher_profile = assignment.teacher
            if teacher_profile.user:
                teacher_info = {
                    "name": teacher_profile.user.full_name or teacher_profile.user.username,
                    "id": teacher_profile.id
                }
            
            # Get primary classroom (first classroom if multiple)
            if teacher_profile.classrooms:
                classroom = teacher_profile.classrooms[0]
                classroom_info = {
                    "name": classroom.class_name,
                    "id": classroom.id
                }
        
        # Get all essays for this assignment
        essays = db.session.query(Essay)\
            .filter(Essay.assignment_id == assignment_id)\
            .all()
        
        logger.info(f"Found {len(essays)} essays for assignment {assignment_id}")
        
        # Build student VMs
        students = []
        for essay in essays:
            student_vm = build_student_vm(essay.id, require_review)
            if student_vm:
                students.append(student_vm)
        
        logger.info(f"Built {len(students)} student VMs for assignment {assignment_id}")
        
        return AssignmentReportVM(
            assignment_id=assignment_id,
            title=assignment.title,
            classroom=classroom_info,
            teacher=teacher_info,
            students=students
        )
        
    except Exception as e:
        logger.error(f"Failed to build assignment VM for assignment {assignment_id}: {e}")
        return None


def build_teacher_view_evaluation(essay_id: int) -> Optional[EvaluationResult]:
    """
    Build EvaluationResult aligned with teacher view data structure.
    
    Args:
        essay_id: Essay ID
        
    Returns:
        EvaluationResult with teacher-view aligned fields populated
    """
    try:
        # Get essay with related data
        essay = db.session.get(Essay, essay_id)
        if not essay:
            logger.warning(f"Essay {essay_id} not found")
            return None
        
        assignment = essay.assignment
        if not assignment:
            logger.warning(f"Assignment not found for essay {essay_id}")
            return None
        
        # Get teacher feedback data (this is what's shown on review page)
        grading_result = essay.ai_score or {}
        if essay.teacher_feedback_overrides:
            try:
                from dictdiffer import patch
                grading_result = patch(essay.teacher_feedback_overrides, grading_result)
            except Exception as e:
                logger.warning(f"Failed to apply teacher feedback overrides for essay {essay_id}: {e}")
                grading_result = essay.ai_score or {}
        
        # Build basic meta information aligned with teacher view
        student_name = "未知学生"
        if essay.enrollment and essay.enrollment.student and essay.enrollment.student.user:
            student_name = essay.enrollment.student.user.full_name or essay.enrollment.student.user.username
        
        submitted_at = essay.created_at.strftime('%Y-%m-%d %H:%M:%S') if essay.created_at else "未知时间"
        
        meta = Meta(
            student=student_name,
            topic=assignment.title,
            date=submitted_at,
            class_="未知班级",  # Could be enhanced to get actual class
            teacher="未知教师"   # Could be enhanced to get actual teacher
        )
        
        # Build scores from grading_result 
        total_score = grading_result.get('total_score', 0)
        rubrics = []
        
        # Map dimensions to rubrics with enhanced data
        dimensions = grading_result.get('dimensions', [])
        for dim in dimensions:
            rubric = RubricScore(
                name=dim.get('dimension_name', ''),
                score=dim.get('score', 0),
                max=dim.get('max_score', 0) or 10,  # Default max if not provided
                weight=1.0,
                reason=dim.get('feedback', '') or dim.get('selected_rubric_level', '')
            )
            
            # Add example sentences data if available
            if 'example_good_sentence' in dim:
                rubric.example_good_sentence = dim['example_good_sentence']
            if 'example_improvement_suggestion' in dim:
                rubric.example_improvement_suggestion = dim['example_improvement_suggestion']
            
            rubrics.append(rubric)
        
        scores = Scores(total=total_score, rubrics=rubrics)
        
        # Get current essay content (teacher's final version)
        # This comes from content_source in the review page
        original_content = essay.content or essay.original_ocr_text or ''
        current_content = essay.teacher_corrected_text if essay.teacher_corrected_text else original_content
        
        text_block = TextBlock(original=original_content, cleaned=current_content)
        
        # Load enhanced evaluation data if available
        evaluation_data = None
        try:
            evaluation_data = load_evaluation_from_essay(essay_id)
        except Exception as e:
            logger.warning(f"Failed to load evaluation data for essay {essay_id}: {e}")
        
        # Extract AI enhanced content from evaluation_data
        outline = []
        diagnoses = []
        personalized_practices = []
        summary_data = None
        parent_summary = ""
        
        if evaluation_data:
            # Map outline from analysis
            if hasattr(evaluation_data, 'analysis') and evaluation_data.analysis:
                outline_items = evaluation_data.analysis.outline if hasattr(evaluation_data.analysis, 'outline') else []
                outline = [{'index': item.para, 'intention': item.intent} for item in outline_items]
            
            # Map diagnostics to diagnoses
            if hasattr(evaluation_data, 'diagnostics') and evaluation_data.diagnostics:
                for i, diag in enumerate(evaluation_data.diagnostics):
                    diagnosis = {
                        'id': i + 1,
                        'target': f"第{diag.para}段" if diag.para else "全文",
                        'evidence': diag.evidence,
                        'suggestions': diag.advice if hasattr(diag, 'advice') else []
                    }
                    diagnoses.append(diagnosis)
            
            # Map exercises to personalized practices
            if hasattr(evaluation_data, 'exercises') and evaluation_data.exercises:
                for ex in evaluation_data.exercises:
                    practice = {
                        'title': ex.type if hasattr(ex, 'type') else '',
                        'requirement': ex.prompt if hasattr(ex, 'prompt') else ''
                    }
                    personalized_practices.append(practice)
            
            # Use summary for parent summary
            if hasattr(evaluation_data, 'summary'):
                parent_summary = evaluation_data.summary
        
        # Create default summary data if not available
        if not summary_data:
            summary_data = {
                'problemSummary': '本次作文分析发现的主要问题包括结构组织、语言表达等方面。',
                'improvementPlan': '建议从基础写作技巧、段落结构、词汇运用等方面进行针对性改进。',
                'expectedOutcome': '通过有针对性的练习和指导，预期能够在作文质量上取得明显提升。'
            }
        
        # Extract meaningful overall comment, strengths, and improvements from dimension feedback
        overall_comment = grading_result.get('overall_comment', '')
        strengths = list(grading_result.get('strengths', []))
        improvements = list(grading_result.get('improvements', []))
        
        # Extract insights from dimension feedback if the dedicated fields are empty
        if not overall_comment and not strengths and not improvements:
            for dim in dimensions:
                feedback = dim.get('feedback', '') or dim.get('selected_rubric_level', '')
                score = dim.get('score', 0)
                max_score = dim.get('max_score', 10)
                score_ratio = score / max_score if max_score > 0 else 0
                
                # If score is high (>=0.8), treat as strength
                if score_ratio >= 0.8 and feedback:
                    strengths.append(f"{dim.get('dimension_name', '')}方面表现优秀：{feedback}")
                # If score is low (<=0.6), treat as improvement area
                elif score_ratio <= 0.6 and feedback:
                    improvements.append(f"{dim.get('dimension_name', '')}需要改进：{feedback}")
            
            # Generate overall comment if still empty
            if not overall_comment:
                total_score = grading_result.get('total_score', 0)
                if total_score >= 32:  # Assuming 40 is max, 80% threshold
                    overall_comment = f"本次作文总体表现良好，获得{total_score}分，显示了扎实的写作基础和良好的表达能力。"
                elif total_score >= 24:  # 60% threshold
                    overall_comment = f"本次作文表现中等，获得{total_score}分，在某些方面表现出色，但仍有进一步提升的空间。"
                else:
                    overall_comment = f"本次作文获得{total_score}分，需要在多个方面加强练习，建议重点关注写作基础技能的提升。"
        
        # Add generic strengths/improvements if still empty
        if not strengths:
            strengths = ["能够完成作文基本要求", "语言表达基本流畅", "内容结构相对完整"]
        
        if not improvements:
            improvements = ["可以进一步丰富内容深度", "语言表达可以更加精准", "文章结构可以更加紧密"]
        
        # Build the EvaluationResult with teacher-view aligned fields
        result = EvaluationResult(
            meta=meta,
            text=text_block,
            scores=scores,
            highlights=[],  # Not needed for current export
            diagnosis=None,  # Using new format instead
            
            # Teacher-view aligned export fields
            assignmentTitle=assignment.title,
            studentName=student_name,
            submittedAt=submitted_at,
            currentEssayContent=current_content,
            
            # AI enhanced content
            outline=outline,
            diagnoses=diagnoses,
            personalizedPractices=personalized_practices,
            summaryData=summary_data,
            parentSummary=parent_summary,
            
            # Additional fields needed for template compatibility with improved data
            overall_comment=overall_comment,
            strengths=strengths,
            improvements=improvements
        )
        
        # Preserve original dimensions data and essay instance for template access
        result._original_grading_result = grading_result
        result._essay_instance = essay  # Pass essay instance for image access
        
        return result
        
    except Exception as e:
        logger.error(f"Failed to build teacher view evaluation for essay {essay_id}: {e}")
        return None


def render_student_docx(essay_id: int) -> bytes:
    """
    Render single student DOCX report using legacy format.
    
    Args:
        essay_id: Essay ID
        
    Returns:
        DOCX file content as bytes
    """
    from app.dao.evaluation_dao import load_evaluation_by_essay
    
    evaluation = load_evaluation_by_essay(essay_id)
    if not evaluation:
        raise ValueError(f"No evaluation found for essay {essay_id}")
    
    # Use existing renderer with legacy format
    output_path = render_essay_docx(evaluation, teacher_view=False)
    
    # Read file and return bytes
    with open(output_path, 'rb') as f:
        return f.read()


def render_teacher_view_docx(essay_id: int) -> bytes:
    """
    Render DOCX report aligned with teacher view (no diff).
    
    Args:
        essay_id: Essay ID
        
    Returns:
        DOCX file content as bytes
        
    Raises:
        ValueError: If evaluation data not found
    """
    evaluation = build_teacher_view_evaluation(essay_id)
    if not evaluation:
        raise ValueError(f"No evaluation data found for essay {essay_id}")
    
    # Use existing renderer with teacher-view aligned evaluation
    output_path = render_essay_docx(evaluation)
    
    # Read file and return bytes
    with open(output_path, 'rb') as f:
        return f.read()
    """
    Render single student DOCX report.
    
    Args:
        essay_id: Essay ID
        
    Returns:
        DOCX file content as bytes
    """
    from app.dao.evaluation_dao import load_evaluation_by_essay
    
    evaluation = load_evaluation_by_essay(essay_id)
    if not evaluation:
        raise ValueError(f"No evaluation found for essay {essay_id}")
    
    # Use existing renderer
    output_path = render_essay_docx(evaluation)
    
    # Read file and return bytes
    with open(output_path, 'rb') as f:
        return f.read()


def render_assignment_docx(assignment_id: int, mode: Literal["combined", "zip"] = "combined", require_review: bool = None) -> Union[bytes, zipstream.ZipStream]:
    """
    Render assignment batch DOCX report.
    
    Args:
        assignment_id: Assignment ID
        mode: "combined" for single merged DOCX, "zip" for ZIP of individual files
        require_review: Whether to require teacher review before export.
                       If None, uses config EVAL_REQUIRE_REVIEW_BEFORE_EXPORT
        
    Returns:
        Bytes for combined mode, ZipFile for zip mode
        
    Raises:
        ValueError: If require_review is True but some evaluations are not teacher-reviewed
    """
    assignment_vm = build_assignment_vm(assignment_id, require_review)
    if not assignment_vm:
        raise ValueError(f"No data found for assignment {assignment_id}")
    
    if not assignment_vm.students:
        raise ValueError(f"No student data found for assignment {assignment_id}")
    
    if mode == "zip":
        return _render_assignment_zip(assignment_vm)
    else:
        return _render_assignment_combined(assignment_vm)


def render_assignment_docx_teacher_view(assignment_id: int, mode: Literal["combined", "zip"] = "combined", require_review: bool = None) -> Union[bytes, zipstream.ZipStream]:
    """
    Render assignment batch DOCX report using teacher view format (no diff).
    
    Args:
        assignment_id: Assignment ID
        mode: "combined" for single merged DOCX, "zip" for ZIP of individual files
        require_review: Whether to require teacher review before export.
                       If None, uses config EVAL_REQUIRE_REVIEW_BEFORE_EXPORT
        
    Returns:
        Bytes for combined mode, ZipStream for zip mode
        
    Raises:
        ValueError: If require_review is True but some evaluations are not teacher-reviewed
    """
    assignment_vm = build_assignment_vm(assignment_id, require_review)
    if not assignment_vm:
        raise ValueError(f"No data found for assignment {assignment_id}")
    
    if not assignment_vm.students:
        raise ValueError(f"No student data found for assignment {assignment_id}")
    
    if mode == "zip":
        return _render_assignment_zip_teacher_view(assignment_vm)
    else:
        return _render_assignment_combined_teacher_view(assignment_vm)


def _render_assignment_combined(assignment_vm: AssignmentReportVM) -> bytes:
    """
    Render combined DOCX using docxtpl with subdocuments.
    
    Args:
        assignment_vm: Assignment data
        
    Returns:
        Combined DOCX as bytes
    """
    try:
        return _render_with_docxtpl_combined(assignment_vm)
    except (ImportError, FileNotFoundError) as e:
        # Only fallback for missing dependencies or IO issues
        logger.warning(f"docxtpl rendering failed due to missing dependency/file: {e}, falling back to docxcompose")
        return _render_with_docxcompose(assignment_vm)
    except Exception as e:
        # Template and context errors should be raised to help debugging
        logger.error(f"Template rendering failed: {e}")
        raise


def _render_with_docxtpl_combined(assignment_vm: AssignmentReportVM) -> bytes:
    """Render using docxtpl with subdocuments and enhanced templates."""
    from docxtpl import DocxTemplate
    from datetime import datetime
    from app.dao.evaluation_dao import load_evaluation_by_essay
    from app.schemas.evaluation import to_context
    import os
    from pathlib import Path
    
    # Ensure assignment template exists, create if missing
    from app.reporting.docx_renderer import ensure_assignment_template_exists
    assignment_template_path = ensure_assignment_template_exists()
    
    # Load the main assignment template
    doc = DocxTemplate(assignment_template_path)
    
    # Prepare context with enhanced student data

    students_data = []
    for student in assignment_vm.students:
        from app.reporting.image_overlay import compose_annotations

        student_data = {
            'student_name': student.student_name,
            'student_no': student.student_no,
            'topic': student.topic,
            'words': student.words,
            # 在 _render_with_docxtpl_combined 构建 student_data 的 'scores' 时，补充一个 rubrics 别名
            'scores': {
                'total': student.scores.total,
                'items': [
                    {
                        'name': item.name,
                        'score': item.score,
                        'max_score': item.max_score,
                        'weight': getattr(item, 'weight', ''),
                        'reason': getattr(item, 'reason', '')
                    }
                    for item in student.scores.items
                ],
                # 兼容字段，方便模板将来统一用 rubrics/max
                'rubrics': [
                    {
                        'name': item.name,
                        'score': item.score,
                        'max': item.max_score,
                        'weight': getattr(item, 'weight', ''),
                        'reason': getattr(item, 'reason', '')
                    }
                    for item in student.scores.items
                ],
            },
            'text': {
                # 优先 cleaned_text，其次用原文段落拼接，再退到反馈
                'cleaned': (
                    getattr(student, 'cleaned_text', '') or
                    ("\n".join(student.original_paragraphs) if getattr(student, 'original_paragraphs', None) else '') or
                    student.feedback
                )
            },
            'analysis': {
                'outline': getattr(student, 'outline', []),
                'issues': getattr(student, 'issues', [])
            },
            'diagnostics': getattr(student, 'diagnostics', []),
            'diagnosis': {
                'before': getattr(student, 'diagnosis_before', ''),
                'comment': getattr(student, 'diagnosis_comment', ''),
                'after': getattr(student, 'diagnosis_after', '')
            },
            'summary': getattr(student, 'summary', ''),
            'paragraphs': [
                {
                    'para_num': para.para_num,
                    'original_text': para.original_text,
                    'feedback': para.feedback,
                    'polished_text': para.polished_text,
                    'intent': para.intent
                } for para in student.paragraphs
            ],
            'exercises': [
                {
                    'type': ex.type,
                    'prompt': ex.prompt,
                    'hint': ex.hints if hasattr(ex, 'hints') else getattr(ex, 'hint', []),
                    'sample': ex.sample
                } for ex in student.exercises
            ],
            'images': {
                'original_image_path': None,
                'composited_image_path': None
            },
            'feedback_summary': student.feedback_summary
        }

        # 如缺失则从 EvaluationResult 做兜底回填
        if (not student_data['analysis']['outline'] and not student_data['analysis']['issues']) or not student_data['diagnostics']:
            try:
                evaluation = load_evaluation_by_essay(student.essay_id)
                if evaluation:
                    ctx = to_context(evaluation)
                    analysis_ctx = ctx.get('analysis', {}) or {}
                    student_data['analysis']['outline'] = student_data['analysis']['outline'] or analysis_ctx.get('outline', [])
                    student_data['analysis']['issues'] = student_data['analysis']['issues'] or analysis_ctx.get('issues', [])
                    student_data['diagnostics'] = student_data['diagnostics'] or ctx.get('diagnostics', []) or []
                    diag = ctx.get('diagnosis', {}) or {}
                    for k in ('before', 'comment', 'after'):
                        if not student_data['diagnosis'].get(k):
                            student_data['diagnosis'][k] = diag.get(k, '')
            except Exception:
                # 兜底失败忽略，不影响整体报告
                pass

        # 图片仍需上线游提供 scanned_images 才能显示
        if hasattr(student, 'scanned_images') and student.scanned_images:
            original_path = student.scanned_images[0]
            student_data['images']['original_image_path'] = original_path
            annotations = getattr(student, 'annotations', None)
            if original_path and annotations:
                composited_path = compose_annotations(original_path, annotations)
                student_data['images']['composited_image_path'] = composited_path

        students_data.append(student_data)

    context = {
        'assignment': {
            'title': assignment_vm.title,
            'classroom': assignment_vm.classroom,
            'teacher': assignment_vm.teacher
        },
        'students': students_data,
        'current_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'now': datetime.now()
    }
    # Debug: print context structure
    logger.info(f"Rendering assignment with {len(students_data)} students")
    
    # Register strftime filter for Jinja2 environment
    from jinja2 import Environment
    env = Environment(autoescape=False)
    
    def strftime_filter(dt, fmt):
        """Custom strftime filter that handles both datetime objects and strings"""
        if dt is None:
            return ''
        if isinstance(dt, str):
            return dt  # If already a string, return as-is
        if hasattr(dt, 'strftime'):
            return dt.strftime(fmt)
        return str(dt)
    
    env.filters['strftime'] = strftime_filter
    
    # Render template
    try:
        doc.render(context, jinja_env=env)
    except Exception as e:
        logger.error(f"Template rendering failed: {e}")
        logger.error(f"Context keys: {list(context.keys())}")
        logger.error(f"Students data type: {type(context['students'])}")
        if context['students']:
            logger.error(f"First student keys: {list(context['students'][0].keys())}")
        raise
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _render_with_docxtpl_basic(assignment_vm: AssignmentReportVM, template_path: str) -> bytes:
    """Fallback rendering using basic template (original implementation)."""
    from docxtpl import DocxTemplate
    from datetime import datetime
    
    # Create a combined context with all student data
    combined_context = {
        'assignment': {
            'title': assignment_vm.title,
            'classroom': assignment_vm.classroom,
            'teacher': assignment_vm.teacher
        },
        'students': []
    }
    
    # Add student contexts
    for student in assignment_vm.students:
        from app.dao.evaluation_dao import load_evaluation_by_essay
        evaluation = load_evaluation_by_essay(student.essay_id)
        if evaluation:
            from app.schemas.evaluation import to_context
            student_context = to_context(evaluation)
            student_context['student_name'] = student.student_name
            combined_context['students'].append(student_context)
    
    # Register strftime filter for Jinja2 environment
    from jinja2 import Environment
    env = Environment(autoescape=False)
    
    def strftime_filter(dt, fmt):
        """Custom strftime filter that handles both datetime objects and strings"""
        if dt is None:
            return ''
        if isinstance(dt, str):
            return dt  # If already a string, return as-is
        if hasattr(dt, 'strftime'):
            return dt.strftime(fmt)
        return str(dt)
    
    env.filters['strftime'] = strftime_filter
    
    # Render template
    doc = DocxTemplate(template_path)
    
    # For now, render first student as representative
    if combined_context['students']:
        context = combined_context['students'][0]
        context['now'] = datetime.now()  # Add now for strftime filter
        doc.render(context, jinja_env=env)
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _render_with_docxcompose(assignment_vm: AssignmentReportVM) -> bytes:
    """Render using docxcompose to combine individual documents."""
    from docx import Document
    from docxcompose.composer import Composer
    
    # Generate individual documents
    temp_files = []
    try:
        for i, student in enumerate(assignment_vm.students):
            student_bytes = render_student_docx(student.essay_id)
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_file.write(student_bytes)
            temp_file.close()
            temp_files.append(temp_file.name)
        
        if not temp_files:
            raise ValueError("No documents to combine")
        
        # Compose documents
        master = Document(temp_files[0])
        composer = Composer(master)
        
        for path in temp_files[1:]:
            composer.append(Document(path))
        
        # Save to bytes
        output = io.BytesIO()
        master.save(output)
        return output.getvalue()
        
    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                Path(temp_file).unlink()
            except:
                pass


def _render_assignment_zip(assignment_vm: AssignmentReportVM) -> zipstream.ZipStream:
    """
    Render assignment as ZIP of individual DOCX files.
    
    Args:
        assignment_vm: Assignment data
        
    Returns:
        ZipStream generator
    """
    z = zipstream.ZipStream(compress_type=zipstream.ZIP_DEFLATED)
    
    for student in assignment_vm.students:
        try:
            student_bytes = render_student_docx(student.essay_id)
            filename = f"{student.student_name}_{student.topic}_{student.essay_id}.docx"
            # Sanitize filename
            filename = "".join(c for c in filename if c.isalnum() or c in "._-")
            z.add(student_bytes, filename)
        except Exception as e:
            logger.error(f"Failed to render student {student.student_name}: {e}")
            continue
    
    return z


def _render_assignment_zip_teacher_view(assignment_vm: AssignmentReportVM) -> zipstream.ZipStream:
    """
    Render assignment as ZIP of individual teacher view DOCX files.
    
    Args:
        assignment_vm: Assignment data
        
    Returns:
        ZipStream generator
        
    Raises:
        ValueError: If no documents could be generated
    """
    z = zipstream.ZipStream(compress_type=zipstream.ZIP_DEFLATED)
    
    successful_count = 0
    failed_students = []
    
    for student in assignment_vm.students:
        try:
            # Use teacher view instead of legacy student format
            student_bytes = render_teacher_view_docx(student.essay_id)
            # Use assignment title and timestamp for better naming
            assignment_title = assignment_vm.title or "assignment"
            safe_assignment_title = "".join(c for c in assignment_title if c.isalnum() or c in "._-")
            safe_student_name = "".join(c for c in student.student_name if c.isalnum() or c in "._-") 
            from datetime import datetime
            timestamp = datetime.now().strftime('%Y%m%d')
            filename = f"{safe_student_name}_{safe_assignment_title}_{timestamp}.docx"
            z.add(student_bytes, filename)
            successful_count += 1
        except Exception as e:
            logger.error(f"Failed to render teacher view for student {student.student_name} (essay_id: {student.essay_id}): {e}")
            failed_students.append({
                'student_name': student.student_name,
                'essay_id': student.essay_id,
                'error': str(e)
            })
            continue
    
    # Add validation consistent with combined mode
    if successful_count == 0:
        total_students = len(assignment_vm.students)
        error_details = []
        for failure in failed_students:
            error_details.append(f"- {failure['student_name']} (essay_id: {failure['essay_id']}): {failure['error']}")
        
        detailed_error = (
            f"No teacher view documents could be generated for assignment '{assignment_vm.title}'. "
            f"All {total_students} students failed:\n" + "\n".join(error_details) + 
            f"\n\nCommon causes:\n"
            f"- Essays have no evaluation data (check if AI grading completed)\n"
            f"- Essays are missing from database\n" 
            f"- Teacher review required but not completed\n"
            f"- Data integrity issues with essay/evaluation records"
        )
        raise ValueError(detailed_error)
    
    # Log summary of results
    if failed_students:
        logger.warning(f"ZIP export partially successful: {successful_count}/{len(assignment_vm.students)} students. "
                      f"Failed students: {[f['student_name'] for f in failed_students]}")
    else:
        logger.info(f"ZIP export successful for all {successful_count} students")
    
    return z


def _render_assignment_combined_teacher_view(assignment_vm: AssignmentReportVM) -> bytes:
    """
    Render combined DOCX using docxcompose with individual teacher view DOCX files.
    
    Args:
        assignment_vm: Assignment data
        
    Returns:
        Combined DOCX as bytes
    """
    from docx import Document
    from docxcompose.composer import Composer
    import tempfile
    from pathlib import Path
    
    # Generate individual teacher view documents
    temp_files = []
    failed_students = []
    try:
        for i, student in enumerate(assignment_vm.students):
            try:
                student_bytes = render_teacher_view_docx(student.essay_id)
                
                # Save to temporary file
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                temp_file.write(student_bytes)
                temp_file.close()
                temp_files.append(temp_file.name)
            except Exception as e:
                logger.error(f"Failed to render teacher view for student {student.student_name} (essay_id: {student.essay_id}): {e}")
                failed_students.append({
                    'student_name': student.student_name,
                    'essay_id': student.essay_id,
                    'error': str(e)
                })
                continue
        
        if not temp_files:
            # Provide detailed error information for debugging
            total_students = len(assignment_vm.students)
            error_details = []
            for failure in failed_students:
                error_details.append(f"- {failure['student_name']} (essay_id: {failure['essay_id']}): {failure['error']}")
            
            detailed_error = (
                f"No teacher view documents could be generated for assignment '{assignment_vm.title}'. "
                f"All {total_students} students failed:\n" + "\n".join(error_details) + 
                f"\n\nCommon causes:\n"
                f"- Essays have no evaluation data (check if AI grading completed)\n"
                f"- Essays are missing from database\n"
                f"- Teacher review required but not completed\n"
                f"- Data integrity issues with essay/evaluation records"
            )
            raise ValueError(detailed_error)
        
        # Log summary of results
        successful_count = len(temp_files)
        total_students = len(assignment_vm.students)
        if failed_students:
            logger.warning(f"Combined export partially successful: {successful_count}/{total_students} students. "
                          f"Failed students: {[f['student_name'] for f in failed_students]}")
        else:
            logger.info(f"Combined export successful for all {successful_count} students")
        
        # Create cover page and compose documents
        master = Document(temp_files[0])
        
        # Insert assignment info at the beginning
        cover_para = master.paragraphs[0]
        cover_para.insert_paragraph_before(f"作业批量报告")
        cover_para.insert_paragraph_before(f"作业标题：{assignment_vm.title}")
        cover_para.insert_paragraph_before(f"班级：{assignment_vm.classroom}")
        cover_para.insert_paragraph_before(f"教师：{assignment_vm.teacher}")
        from datetime import datetime
        cover_para.insert_paragraph_before(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        cover_para.insert_paragraph_before("")
        
        if len(temp_files) > 1:
            composer = Composer(master)
            
            for path in temp_files[1:]:
                # Add page break before each new student
                composer.append(Document(path))
        
        # Save to bytes
        output = io.BytesIO()
        master.save(output)
        return output.getvalue()
        
    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                Path(temp_file).unlink()
            except:
                pass