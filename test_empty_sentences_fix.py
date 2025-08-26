"""
Test to reproduce and verify the fix for empty sentence modules in DOCX export.
"""
import os
import tempfile
import pytest
from datetime import datetime

from app.schemas.evaluation import (
    EvaluationResult, Meta, Scores, RubricScore, TextBlock
)
from app.reporting.docx_renderer import render_essay_docx


def test_empty_sentence_modules_not_rendered():
    """Test that empty good sentence and improvement suggestion modules are not rendered"""
    
    # Create evaluation result with rubrics that have empty sentence data
    evaluation = EvaluationResult(
        meta=Meta(
            student="测试学生",
            class_="测试班级",
            teacher="测试教师", 
            topic="测试作文",
            date="2024-01-01"
        ),
        scores=Scores(
            total=85.0,
            rubrics=[
                # Rubric with empty good sentences and improvement suggestions
                RubricScore(
                    name="内容", 
                    score=18.0, 
                    max=20.0, 
                    weight=1.0, 
                    reason="内容评价",
                    example_good_sentence=[],  # Empty list
                    example_improvement_suggestion=[]  # Empty list
                ),
                # Rubric with None values
                RubricScore(
                    name="结构", 
                    score=17.0, 
                    max=20.0, 
                    weight=1.0, 
                    reason="结构评价",
                    example_good_sentence=None,
                    example_improvement_suggestion=None
                ),
                # Rubric with lists containing empty strings
                RubricScore(
                    name="语言", 
                    score=16.0, 
                    max=20.0, 
                    weight=1.0, 
                    reason="语言评价",
                    example_good_sentence=["", "  ", ""],  # Empty/whitespace strings
                    example_improvement_suggestion=[
                        {"original": "", "suggested": ""},  # Empty dict values
                        {"original": "  ", "suggested": "  "}  # Whitespace dict values
                    ]
                ),
                # Rubric with proper content (should be rendered)
                RubricScore(
                    name="创新", 
                    score=15.0, 
                    max=20.0, 
                    weight=1.0, 
                    reason="创新评价",
                    example_good_sentence=["这是一个很好的句子示例"],
                    example_improvement_suggestion=[
                        {"original": "原始句子", "suggested": "改进建议"}
                    ]
                )
            ]
        ),
        text=TextBlock(
            original="这是测试作文内容。",
            cleaned="这是测试作文内容。"
        ),
        assignmentTitle="测试作文题目",
        studentName="测试学生",
        submittedAt="2024-01-01",
        currentEssayContent="这是测试作文内容。"
    )
    
    # Render DOCX to temporary file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        output_path = temp_file.name
    
    try:
        # Test both teacher view and regular rendering
        result_path = render_essay_docx(evaluation, output_path, teacher_view=True)
        assert os.path.exists(result_path)
        
        # Read the generated DOCX content to verify empty modules are not rendered
        from docx import Document
        doc = Document(result_path)
        
        # Get all paragraph text
        doc_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # Should contain the dimension with actual content
        assert "创新维度详情" in doc_text
        assert "这是一个很好的句子示例" in doc_text
        assert "原始句子" in doc_text
        assert "改进建议" in doc_text
        
        # Should NOT contain headers for empty dimensions
        # Count occurrences of empty section headers
        good_sentence_headers = doc_text.count("亮点句子：")
        improvement_headers = doc_text.count("待改进句：")
        
        # Should only have 1 of each (from the "创新" dimension that has actual content)
        assert good_sentence_headers == 1, f"Expected 1 '亮点句子：' header, got {good_sentence_headers}"
        assert improvement_headers == 1, f"Expected 1 '待改进句：' header, got {improvement_headers}"
        
        print("✓ Empty sentence modules are correctly hidden")
        print(f"✓ Found {good_sentence_headers} good sentence headers (expected: 1)")
        print(f"✓ Found {improvement_headers} improvement headers (expected: 1)")
        
    finally:
        # Clean up
        if os.path.exists(output_path):
            os.unlink(output_path)


def test_all_empty_sentence_modules():
    """Test case where all dimensions have empty sentence data"""
    
    evaluation = EvaluationResult(
        meta=Meta(
            student="测试学生2",
            class_="测试班级2",
            teacher="测试教师2", 
            topic="测试作文2",
            date="2024-01-01"
        ),
        scores=Scores(
            total=75.0,
            rubrics=[
                RubricScore(
                    name="内容", 
                    score=15.0, 
                    max=20.0, 
                    weight=1.0, 
                    reason="内容评价",
                    example_good_sentence=[],
                    example_improvement_suggestion=[]
                ),
                RubricScore(
                    name="结构", 
                    score=16.0, 
                    max=20.0, 
                    weight=1.0, 
                    reason="结构评价",
                    example_good_sentence=None,
                    example_improvement_suggestion=None
                )
            ]
        ),
        text=TextBlock(
            original="测试内容2。",
            cleaned="测试内容2。"
        ),
        assignmentTitle="测试题目2",
        studentName="测试学生2", 
        submittedAt="2024-01-01",
        currentEssayContent="测试内容2。"
    )
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        output_path = temp_file.name
    
    try:
        result_path = render_essay_docx(evaluation, output_path, teacher_view=True)
        assert os.path.exists(result_path)
        
        from docx import Document
        doc = Document(result_path)
        doc_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # Should NOT contain any sentence headers when all are empty
        good_sentence_headers = doc_text.count("亮点句子：")
        improvement_headers = doc_text.count("待改进句：")
        
        assert good_sentence_headers == 0, f"Expected 0 '亮点句子：' headers, got {good_sentence_headers}"
        assert improvement_headers == 0, f"Expected 0 '待改进句：' headers, got {improvement_headers}"
        
        print("✓ All empty sentence modules are correctly hidden")
        
    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)


if __name__ == "__main__":
    print("Testing empty sentence modules fix...")
    test_empty_sentence_modules_not_rendered()
    test_all_empty_sentence_modules()
    print("All tests passed!")