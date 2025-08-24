#!/usr/bin/env python3
"""
Test script to demonstrate the problem with hardcoded fallback data
and verify that real AI evaluation data is used when available.
"""

import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from app.schemas.evaluation import EvaluationResult, Meta, Scores, RubricScore, TextBlock
from app.reporting.docx_renderer import render_essay_docx
import tempfile
import json

def test_with_real_ai_data():
    """Test with actual AI evaluation data structure similar to the database sample"""
    
    # Simulate real AI data structure from database (based on problem statement sample)
    ai_evaluation_data = {
        "meta": {
            "student": "蔡桢达",
            "class_": "",
            "teacher": "沈颖",
            "topic": "写读后感--实用-学术类设计",
            "date": "2025-08-23",
            "student_id": "47",
            "grade": "",
            "words": 0
        },
        "scores": {
            "total": 32,
            "rubrics": [
                {
                    "name": "书籍内容概述提炼",
                    "score": 17,
                    "max": 20,
                    "weight": 1.0,
                    "reason": "在这一维度上表现优秀！你不仅准确概述了《魔道祖师》的基本信息，还特别聚焦于自己最感兴趣的部分——观音庙情节进行了详细描述。",
                    "level": "A",
                    "example_good_sentence": ["《魔道祖师》是墨香铜臭所著的长篇修真小说，主要讲述了夷陵老祖魏无羡献舍归来后，与姑苏蓝氏二公子蓝忘机共同探查莫家庄凶尸奇案的故事。"],
                    "example_improvement_suggestion": []
                },
                {
                    "name": "心得体会交流分享",
                    "score": 8,
                    "max": 10,
                    "weight": 1.0,
                    "reason": "心得体会部分做得很好！你不仅从书中提炼出了'祸从口出'的核心道理，还能结合自己的真实经历来印证这个道理。",
                    "level": "A",
                    "example_good_sentence": ["记得有一次，我和一个朋友开了一个玩笑，说他有点矮，结果出去玩时他拿着篮球砸向我的头。"],
                    "example_improvement_suggestion": []
                },
                {
                    "name": "语言表达学术规范",
                    "score": 7,
                    "max": 10,
                    "weight": 1.0,
                    "reason": "语言表达基本规范，但还有一些提升空间。整体语言流畅，能够准确使用书名号等标点符号。",
                    "level": "B",
                    "example_good_sentence": ["《魔道祖师》是墨香铜臭所著的长篇修真小说，主要讲述了夷陵老祖魏无羡献舍归来后，与姑苏蓝氏二公子蓝忘机共同探查莫家庄凶尸奇案的故事。"],
                    "example_improvement_suggestion": ["优秀的文学作品往往能够给予读者深刻的人生启示。"]
                }
            ]
        },
        "overall_comment": "这篇读后感展现了不错的阅读理解和感悟能力！你能够清晰地概述《魔道祖师》的主要内容和人物特点，特别是对观音庙情节的描写很具体，并且能够从中提炼出'祸从口出'的道理，还能结合自己的生活经历来谈体会，这一点非常值得肯定。整体结构完整，思路清晰，是一篇有思考、有感悟的读后感。",
        "strengths": [
            "能够准确概括书籍主要内容并聚焦自己感兴趣的情节进行详细描述",
            "能够将阅读感悟与自身生活经历相结合，体现了真实的阅读收获"
        ],
        "improvements": [
            "可以进一步深化对书中主题和人物命运的理解，挖掘更深层的启示",
            "语言表达可以更加学术化和规范化，减少口语化表述"
        ],
        "text": {
            "original": "读《魔道祖师》有感\n最近，我读了作者墨香铜臭的《魔道祖师》，颇有感触。\n《魔道祖师》是墨香铜臭所著的长篇修真小说，主要讲述了夷陵老祖魏无羡献舍归来后，与姑苏蓝氏二公子蓝忘机共同探查莫家庄凶尸奇案的故事。",
            "cleaned": "读《魔道祖师》有感\n最近，我读了作者墨香铜臭的《魔道祖师》，颇有感触。\n《魔道祖师》是墨香铜臭所著的长篇修真小说，主要讲述了夷陵老祖魏无羡献舍归来后，与姑苏蓝氏二公子蓝忘机共同探查莫家庄凶尸奇案的故事。"
        },
        "assignmentTitle": "读后感写作练习", 
        "studentName": "蔡桢达",
        "submittedAt": "2025-08-23",
        "currentEssayContent": "读《魔道祖师》有感\n最近，我读了作者墨香铜臭的《魔道祖师》，颇有感触。\n《魔道祖师》是墨香铜臭所著的长篇修真小说，主要讲述了夷陵老祖魏无羡献舍归来后，与姑苏蓝氏二公子蓝忘机共同探查莫家庄凶尸奇案的故事。"
    }
    
    # Create EvaluationResult with real data 
    evaluation = EvaluationResult.model_validate(ai_evaluation_data)
    
    print("=== Testing with Real AI Data ===")
    print(f"Student: {evaluation.meta.student}")
    print(f"Total Score: {evaluation.scores.total}")
    print(f"Number of rubrics: {len(evaluation.scores.rubrics)}")
    print(f"Strengths count: {len(evaluation.strengths)}")
    print(f"Improvements count: {len(evaluation.improvements)}")
    print(f"Actual strengths: {evaluation.strengths}")
    print(f"Actual improvements: {evaluation.improvements}")
    print(f"getattr strengths: {getattr(evaluation, 'strengths', [])}")
    print(f"getattr improvements: {getattr(evaluation, 'improvements', [])}")
    
    # Check if rubric data is properly populated
    for i, rubric in enumerate(evaluation.scores.rubrics):
        print(f"\nRubric {i+1}: {rubric.name}")
        print(f"  Score: {rubric.score}/{rubric.max}")
        print(f"  Has example_good_sentence: {hasattr(rubric, 'example_good_sentence') and rubric.example_good_sentence}")
        print(f"  Has example_improvement_suggestion: {hasattr(rubric, 'example_improvement_suggestion') and rubric.example_improvement_suggestion}")
        if hasattr(rubric, 'example_good_sentence') and rubric.example_good_sentence:
            print(f"  Good sentences: {rubric.example_good_sentence}")
        if hasattr(rubric, 'example_improvement_suggestion') and rubric.example_improvement_suggestion:
            print(f"  Improvements: {rubric.example_improvement_suggestion}")
    
    # Generate DOCX
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation)
        print(f"\nGenerated DOCX: {result_path}")
        
        # Read and analyze the content
        from docx import Document
        doc = Document(result_path)
        
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        print("\n=== Generated Content Analysis ===")
        
        # Check if real strengths are used instead of fallback
        real_strength = "能够准确概括书籍主要内容并聚焦自己感兴趣的情节进行详细描述"
        fallback_strength = "能够完成作文基本要求"
        
        if real_strength in full_text:
            print("✅ GOOD: Real strength data is used")
        elif fallback_strength in full_text:
            print("❌ BAD: Fallback strength data is used instead of real data")
        
        # Check if real improvements are used instead of fallback
        real_improvement = "可以进一步深化对书中主题和人物命运的理解，挖掘更深层的启示"
        fallback_improvement = "建议进一步丰富表达方式，提升语言准确性"
        
        if real_improvement in full_text:
            print("✅ GOOD: Real improvement data is used")
        elif fallback_improvement in full_text:
            print("❌ BAD: Fallback improvement data is used instead of real data")
        
        # Check if real overall comment is used
        real_overall_comment = "这篇读后感展现了不错的阅读理解和感悟能力"
        fallback_overall_comment = "本次作文总体表现良好"
        
        if real_overall_comment in full_text:
            print("✅ GOOD: Real overall comment is used")
        elif fallback_overall_comment in full_text:
            print("❌ BAD: Fallback overall comment is used instead of real data")
        else:
            print("⚠️  UNKNOWN: Neither real nor fallback overall comment found")
            
        return full_text


def test_with_missing_data():
    """Test with missing data to verify fallbacks only apply when data is truly missing"""
    
    # Create minimal evaluation with no AI enhancement data
    evaluation = EvaluationResult(
        meta=Meta(
            student="测试学生",
            class_="测试班级", 
            teacher="测试老师",
            topic="测试作文",
            date="2024-08-21"
        ),
        scores=Scores(
            total=90.0,
            rubrics=[
                RubricScore(name="测试维度", score=18.0, max=20.0, weight=1.0, reason="测试理由")
            ]
        ),
        text=TextBlock(
            original="这是测试文本内容。",
            cleaned="这是测试文本内容。"
        ),
        # Explicitly set empty data
        strengths=[],
        improvements=[]
    )
    
    print("\n\n=== Testing with Missing Data ===") 
    print("This should use fallbacks since no real data is provided")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation)
        
        from docx import Document
        doc = Document(result_path)
        
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # In this case, fallbacks should be used
        fallback_strength = "能够完成作文基本要求"
        fallback_improvement = "建议进一步丰富表达方式，提升语言准确性"
        
        if fallback_strength in full_text:
            print("✅ GOOD: Fallback is used when data is missing")
        else:
            print("❌ BAD: No strength data found")
            
        if fallback_improvement in full_text:
            print("✅ GOOD: Fallback is used when data is missing") 
        else:
            print("❌ BAD: No improvement data found")


if __name__ == "__main__":
    print("Testing DOCX Renderer Data Usage")
    print("=" * 50)
    
    real_data_content = test_with_real_ai_data()
    test_with_missing_data()
    
    print("\n" + "=" * 50)
    print("Test completed. Check the output above to see if real data is being used.")