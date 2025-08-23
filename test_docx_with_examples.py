#!/usr/bin/env python3
"""
Test DOCX generation with dimension examples to verify the fix.
This demonstrates the difference between EvaluationResult with and without _original_grading_result.
"""

import tempfile
import os
from app.schemas.evaluation import (
    EvaluationResult, Meta, Scores, RubricScore, TextBlock
)
from app.reporting.docx_renderer import render_essay_docx


def create_test_evaluation(with_original_grading_result=False):
    """Create a test EvaluationResult with optional original grading result"""
    
    evaluation = EvaluationResult(
        meta=Meta(
            student="张三",
            class_="五年级1班",
            teacher="李老师",
            topic="读书心得",
            date="2024-08-21"
        ),
        scores=Scores(
            total=85.0,
            rubrics=[
                RubricScore(name="书籍内容概述提炼", score=18.0, max=20.0, weight=1.0, reason="内容概述较好"),
                RubricScore(name="心得体会交流分享", score=16.0, max=20.0, weight=1.0, reason="心得体会表达清晰"),
                RubricScore(name="语言表达学术规范", score=15.0, max=20.0, weight=1.0, reason="语言表达需要提升")
            ]
        ),
        text=TextBlock(
            original="这是一篇读书心得的内容。书中的故事很感人。",
            cleaned="这是一篇读书心得的内容。书中的故事很感人。"
        ),
        strengths=['内容丰富', '感情真挚'],
        improvements=['语言表达可以更准确', '可以增加更多例证']
    )
    
    if with_original_grading_result:
        # Attach original grading result with dimension examples
        evaluation._original_grading_result = {
            'total_score': 85,
            'dimensions': [
                {
                    'dimension_name': '书籍内容概述提炼',
                    'score': 18,
                    'feedback': '内容概述较好',
                    'example_good_sentence': '这本书生动地描述了主人公的成长历程，情节跌宕起伏。',
                    'example_improvement_suggestion': {
                        'original': '书很好看。',
                        'suggested': '这本书通过细腻的描写和深刻的主题展现了人生的复杂性。'
                    }
                },
                {
                    'dimension_name': '心得体会交流分享',
                    'score': 16,
                    'feedback': '心得体会表达清晰',
                    'example_good_sentence': '通过阅读这本书，我深深地感受到了友谊的珍贵和生命的意义。',
                    'example_improvement_suggestion': {
                        'original': '我觉得很有道理。',
                        'suggested': '这个观点引发了我对人际关系本质的深入思考。'
                    }
                },
                {
                    'dimension_name': '语言表达学术规范',
                    'score': 15,
                    'feedback': '语言表达需要提升',
                    'example_good_sentence': '作者运用了丰富的修辞手法，使文章生动有趣，引人深思。',
                    'example_improvement_suggestion': {
                        'original': '写得好。',
                        'suggested': '文章结构严谨，论证有力，语言准确生动。'
                    }
                }
            ],
            'strengths': ['内容丰富', '感情真挚'],
            'improvements': ['语言表达可以更准确', '可以增加更多例证']
        }
    
    return evaluation


def test_docx_generation():
    """Test DOCX generation with and without dimension examples"""
    
    print("Testing DOCX generation...")
    
    # Test without original grading result (should show "（本项暂无数据）")
    print("\n1. Testing without dimension examples (before fix):")
    evaluation_without_examples = create_test_evaluation(with_original_grading_result=False)
    
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation_without_examples)
        print(f"   Generated DOCX (without examples): {result_path}")
        
        # Read content to verify it shows "no data"
        try:
            from docx import Document
            doc = Document(result_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            no_data_count = full_text.count("（本项暂无数据）")
            print(f"   Found {no_data_count} instances of '（本项暂无数据）' (expected: multiple)")
            
        except Exception as e:
            print(f"   Could not read DOCX content: {e}")
    
    # Test with original grading result (should show actual examples)
    print("\n2. Testing with dimension examples (after fix):")
    evaluation_with_examples = create_test_evaluation(with_original_grading_result=True)
    
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation_with_examples)
        print(f"   Generated DOCX (with examples): {result_path}")
        
        # Read content to verify examples are included
        try:
            from docx import Document
            doc = Document(result_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            no_data_count = full_text.count("（本项暂无数据）")
            print(f"   Found {no_data_count} instances of '（本项暂无数据）' (expected: fewer)")
            
            # Check for example content
            has_good_sentence = "这本书生动地描述了主人公的成长历程" in full_text
            has_improvement = "这本书通过细腻的描写和深刻的主题" in full_text
            
            print(f"   Contains example good sentence: {has_good_sentence}")
            print(f"   Contains improvement suggestion: {has_improvement}")
            
            if has_good_sentence and has_improvement:
                print("   ✅ SUCCESS: Dimension examples are properly included!")
            else:
                print("   ❌ Issue: Dimension examples not found in output")
                
        except Exception as e:
            print(f"   Could not read DOCX content: {e}")
    
    print("\nTest completed. The fix should reduce the number of '（本项暂无数据）' instances.")


if __name__ == "__main__":
    test_docx_generation()