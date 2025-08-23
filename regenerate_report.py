#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
重新生成作业报告的脚本
"""

import sys
import os
import json
from datetime import datetime

# 添加项目根目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app import create_app
from app.extensions import db
from app.models import AssignmentReport, Essay
from app.services.ai_report_analyzer import analyze_assignment_with_ai
from app.schemas.evaluation import EvaluationResult

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def regenerate_assignment_report(assignment_id, user_id=1):
    """重新生成指定作业的报告"""
    app = create_app()
    
    with app.app_context():
        try:
            print(f"开始重新生成作业 {assignment_id} 的报告...")
            
            # 删除现有报告（如果存在）
            existing_report = AssignmentReport.query.filter_by(assignment_id=assignment_id).first()
            if existing_report:
                print("删除现有报告...")
                db.session.delete(existing_report)
                db.session.commit()
            
            # 生成新报告
            print("调用AI分析器生成新报告...")
            report_data = analyze_assignment_with_ai(assignment_id)
            
            # 保存报告到数据库
            print("保存报告到数据库...")
            new_report = AssignmentReport(
                assignment_id=assignment_id,
                report_data=json.dumps(report_data, ensure_ascii=False),
                generated_by=user_id
            )
            
            db.session.add(new_report)
            db.session.commit()
            
            print("报告生成成功！")
            
            # 输出问题类型分布数据用于验证
            if 'charts' in report_data and 'problem_types' in report_data['charts']:
                problem_types = report_data['charts']['problem_types']
                print("\n问题类型分布数据：")
                print(f"标签: {problem_types['labels']}")
                print(f"数据: {problem_types['data']}")
            
            return True
            
        except Exception as e:
            print(f"生成报告失败: {e}")
            db.session.rollback()
            return False


def generate_word_report_from_evaluation(essay_id, output_path=None, app=None):
    """
    从EvaluationResult JSON生成Word文档
    
    Args:
        essay_id: 作文ID
        output_path: 输出路径，如果为None则自动生成
        app: Flask app实例，用于测试
        
    Returns:
        生成的文档路径或None（如果失败）
    """
    if app is None:
        app = create_app()
        needs_context = True
    else:
        needs_context = False
    
    def _do_work():
        try:
            # 获取作文记录
            essay = db.session.get(Essay, essay_id)
            if not essay:
                print(f"作文 {essay_id} 不存在")
                return None
            
            # 检查是否有新格式的评估结果
            if not essay.ai_score:
                print(f"作文 {essay_id} 没有评分数据")
                return None
            
            # 尝试解析为新的EvaluationResult格式
            try:
                evaluation = EvaluationResult.model_validate(essay.ai_score)
                print(f"使用新格式EvaluationResult生成DOCX报告")
            except Exception:
                print(f"作文 {essay_id} 使用旧格式，暂不支持")
                return None
            
            # 生成输出路径
            final_output_path = output_path
            if final_output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"essay_{essay_id}_report_{timestamp}.docx"
                final_output_path = os.path.join(os.path.dirname(__file__), "reports", filename)
                os.makedirs(os.path.dirname(final_output_path), exist_ok=True)
            
            # 确保文件扩展名是.docx
            if not final_output_path.endswith('.docx'):
                final_output_path = final_output_path.rsplit('.', 1)[0] + '.docx'
            
            # 生成DOCX报告
            if DOCX_AVAILABLE:
                result_path = generate_docx_report_from_evaluation(evaluation, essay, final_output_path)
                print(f"DOCX报告已生成: {result_path}")
                return result_path
            else:
                # Fallback to text if DOCX not available
                text_path = final_output_path.replace('.docx', '.txt')
                content = generate_report_content(evaluation, essay)
                with open(text_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"DOCX不可用，已生成文本报告: {text_path}")
                return text_path
            
        except Exception as e:
            print(f"生成报告失败: {e}")
            return None
    
    if needs_context:
        with app.app_context():
            return _do_work()
    else:
        return _do_work()


def generate_report_content(evaluation: EvaluationResult, essay):
    """生成报告内容"""
    lines = []
    lines.append("=" * 50)
    lines.append("作文评估报告")
    lines.append("=" * 50)
    lines.append("")
    
    # 基本信息
    lines.append("【基本信息】")
    lines.append(f"学生ID: {evaluation.meta.student_id or '未知'}")
    lines.append(f"年级: {evaluation.meta.grade}")
    lines.append(f"题目: {evaluation.meta.topic or '未知'}")
    lines.append(f"字数: {evaluation.meta.words}")
    lines.append("")
    
    # 评分结果
    lines.append("【评分结果】")
    lines.append(f"总分: {evaluation.scores.total}分")
    lines.append(f"内容: {evaluation.scores.content}分")
    lines.append(f"结构: {evaluation.scores.structure}分") 
    lines.append(f"语言: {evaluation.scores.language}分")
    lines.append(f"文采: {evaluation.scores.aesthetics}分")
    lines.append(f"规范: {evaluation.scores.norms}分")
    lines.append(f"评分理由: {evaluation.scores.rationale}")
    lines.append("")
    
    # 结构分析
    lines.append("【结构分析】")
    for item in evaluation.analysis.outline:
        lines.append(f"第{item.para}段: {item.intent}")
    lines.append("")
    
    # 问题分析
    if evaluation.analysis.issues:
        lines.append("【问题分析】")
        for i, issue in enumerate(evaluation.analysis.issues, 1):
            lines.append(f"{i}. {issue}")
        lines.append("")
    
    # 诊断建议
    if evaluation.diagnostics:
        lines.append("【诊断建议】")
        for diag in evaluation.diagnostics:
            para_info = f"第{diag.para}段" if diag.para else "全文"
            lines.append(f"{para_info} - {diag.issue}")
            lines.append(f"  证据: {diag.evidence}")
            if diag.advice:
                lines.append(f"  建议: {', '.join(diag.advice)}")
        lines.append("")
    
    # 练习建议  
    if evaluation.exercises:
        lines.append("【练习建议】")
        for ex in evaluation.exercises:
            lines.append(f"练习类型: {ex.type}")
            lines.append(f"提示: {ex.prompt}")
            if ex.hint:
                lines.append(f"要点: {', '.join(ex.hint)}")
            if ex.sample:
                lines.append(f"示例: {ex.sample}")
            lines.append("")
    
    # 总结
    if evaluation.summary:
        lines.append("【总结】")
        lines.append(evaluation.summary)
        lines.append("")
    
    lines.append("=" * 50)
    lines.append(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    return '\n'.join(lines)


# 在现有 generate_docx_report_from_evaluation 中，对“诊断建议/练习建议”做轻量版式增强
def _as_bullet(paragraph):
    try:
        paragraph.style = 'List Bullet'
    except Exception:
        pass
    return paragraph

def generate_docx_report_from_evaluation(evaluation: EvaluationResult, essay, output_path: str):
    """Generate DOCX report from EvaluationResult"""
    
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available. Install with: pip install python-docx")
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('作文评估报告', 0)
        title.alignment = 1  # Center alignment
        
        # Basic information
        doc.add_heading('基本信息', level=1)
        basic_info = doc.add_paragraph()
        basic_info.add_run(f"学生ID: ").bold = True
        basic_info.add_run(f"{evaluation.meta.student_id or '未知'}\n")
        basic_info.add_run(f"年级: ").bold = True
        basic_info.add_run(f"{evaluation.meta.grade}\n")
        basic_info.add_run(f"题目: ").bold = True
        basic_info.add_run(f"{evaluation.meta.topic or '未知'}\n")
        basic_info.add_run(f"字数: ").bold = True
        basic_info.add_run(f"{evaluation.meta.words}")
        
        # Scores
        doc.add_heading('评分结果', level=1)
        scores_para = doc.add_paragraph()
        scores_para.add_run(f"总分: ").bold = True
        scores_para.add_run(f"{evaluation.scores.total}分\n")
        scores_para.add_run(f"内容: ").bold = True
        scores_para.add_run(f"{evaluation.scores.content}分\n")
        scores_para.add_run(f"结构: ").bold = True
        scores_para.add_run(f"{evaluation.scores.structure}分\n")
        scores_para.add_run(f"语言: ").bold = True
        scores_para.add_run(f"{evaluation.scores.language}分\n")
        scores_para.add_run(f"文采: ").bold = True
        scores_para.add_run(f"{evaluation.scores.aesthetics}分\n")
        scores_para.add_run(f"规范: ").bold = True
        scores_para.add_run(f"{evaluation.scores.norms}分")
        
        # Rationale
        if evaluation.scores.rationale:
            doc.add_heading('评分说明', level=1)
            doc.add_paragraph(evaluation.scores.rationale)
        
        # Structure analysis
        if evaluation.analysis.outline:
            doc.add_heading('结构分析', level=1)
            for i, item in enumerate(evaluation.analysis.outline, 1):
                para = doc.add_paragraph()
                para.add_run(f"第{item.para}段: ").bold = True
                para.add_run(item.intent)
        
        # Issues
        if evaluation.analysis.issues:
            doc.add_heading('问题分析', level=1)
            for i, issue in enumerate(evaluation.analysis.issues, 1):
                doc.add_paragraph(f"{i}. {issue}")
        
        # Diagnostics
        if evaluation.diagnostics:
            doc.add_heading('诊断建议', level=1)
            for diag in evaluation.diagnostics:
                para_info = f"第{diag.para}段" if diag.para else "全文"
                para = _as_bullet(doc.add_paragraph())
                run = para.add_run(f"{para_info} - {diag.issue}  ")
                run.bold = True

                ev = para.add_run("证据: ")
                ev.bold = True
                para.add_run(f"{diag.evidence}  ")

                if diag.advice:
                    adv = para.add_run("建议: ")
                    adv.bold = True
                    para.add_run('；'.join(diag.advice))

        # Exercises
        if evaluation.exercises:
            doc.add_heading('练习建议', level=1)
            for ex in evaluation.exercises:
                para = _as_bullet(doc.add_paragraph())
                title_run = para.add_run(f"[{ex.type}] {ex.prompt}  ")
                title_run.bold = True
                if getattr(ex, 'hint', None):
                    h = para.add_run("要点: ")
                    h.bold = True
                    para.add_run('；'.join(ex.hint) if isinstance(ex.hint, (list, tuple)) else str(ex.hint))
                if getattr(ex, 'sample', None):
                    para.add_run("  示例: ")
                    para.add_run(ex.sample)

        # Summary
        if evaluation.summary:
            doc.add_heading('总结', level=1)
            doc.add_paragraph(evaluation.summary)
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.alignment = 1  # Center alignment
        
        # Save document
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"生成DOCX报告时发生错误: {e}")
        raise


def _generate_text_report_from_evaluation(essay_id, output_path=None, app=None):
    """Generate text report from EvaluationResult JSON - for backward compatibility"""
    
    if app is None:
        app = create_app()
        needs_context = True
    else:
        needs_context = False
    
    def _do_work():
        try:
            # 获取作文记录
            essay = db.session.get(Essay, essay_id)
            if not essay:
                print(f"作文 {essay_id} 不存在")
                return None
            
            # 检查是否有新格式的评估结果
            if not essay.ai_score:
                print(f"作文 {essay_id} 没有评分数据")
                return None
            
            # 尝试解析为新的EvaluationResult格式
            try:
                evaluation = EvaluationResult.model_validate(essay.ai_score)
                print(f"使用新格式EvaluationResult生成文本报告")
            except Exception:
                print(f"作文 {essay_id} 使用旧格式，暂不支持")
                return None
            
            # 生成输出路径
            final_output_path = output_path
            if final_output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"essay_{essay_id}_report_{timestamp}.txt"
                final_output_path = os.path.join(os.path.dirname(__file__), "reports", filename)
                os.makedirs(os.path.dirname(final_output_path), exist_ok=True)
            
            # 生成报告内容
            content = generate_report_content(evaluation, essay)
            
            # 写入文件
            with open(final_output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            print(f"文本报告已生成: {final_output_path}")
            return final_output_path
            
        except Exception as e:
            print(f"生成文本报告时发生错误: {e}")
            return None
    
    if needs_context:
        with app.app_context():
            return _do_work()
    else:
        return _do_work()


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='重新生成报告')
    parser.add_argument('--assignment', type=int, help='作业ID（生成作业报告）')
    parser.add_argument('--essay-id', type=int, help='作文ID（生成作文报告）')
    parser.add_argument('--format', choices=['txt', 'docx'], default='docx', help='输出格式（默认：docx）')
    parser.add_argument('--output', type=str, help='输出路径')
    
    args = parser.parse_args()
    
    if args.assignment:
        success = regenerate_assignment_report(args.assignment)
        if success:
            print(f"\n作业报告重新生成完成！")
        else:
            print(f"\n作业报告生成失败！")
            sys.exit(1)
    
    elif args.essay_id:
        if args.format == 'docx':
            result = generate_word_report_from_evaluation(args.essay_id, args.output)
        else:
            # Generate text format for backward compatibility - define function below
            result = _generate_text_report_from_evaluation(args.essay_id, args.output)
        
        if result:
            print(f"\n作文报告生成完成: {result}")
        else:
            print(f"\n作文报告生成失败！")
            sys.exit(1)
    
    else:
        # 向后兼容：默认行为
        assignment_id = 2  # 作业ID
        success = regenerate_assignment_report(assignment_id)
        
        if success:
            print("\n报告重新生成完成！")
        else:
            print("\n报告生成失败！")
            sys.exit(1)