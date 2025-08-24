"""
Test specifically for the path resolver functionality in DOCX rendering.
This ensures that Windows paths are properly resolved and friendly messages are shown.
"""
import os
import tempfile
import pytest
from unittest.mock import patch, MagicMock

from app.utils.path_resolver import resolve_upload_path, get_friendly_image_message


def test_path_resolver_handles_windows_paths():
    """Test that path resolver correctly handles Windows absolute paths"""
    
    with tempfile.TemporaryDirectory() as temp_dir:
        # Create uploads structure
        uploads_dir = os.path.join(temp_dir, 'uploads')
        os.makedirs(uploads_dir)
        
        # Create test image
        test_image = os.path.join(uploads_dir, 'image123.jpg')
        with open(test_image, 'w') as f:
            f.write('test image')
        
        # Mock the uploads directory function
        with patch('app.utils.path_resolver._get_uploads_directory', return_value=uploads_dir):
            # Test Windows absolute path resolution
            windows_path = r'D:\Github\evzj\uploads\image123.jpg'
            resolved_path = resolve_upload_path(windows_path)
            
            assert resolved_path == test_image
            assert os.path.exists(resolved_path)
            
            # Test Unix-style path resolution
            unix_path = 'D:/Github/evzj/uploads/image123.jpg'
            resolved_path = resolve_upload_path(unix_path)
            
            assert resolved_path == test_image
            assert os.path.exists(resolved_path)


def test_path_resolver_returns_none_for_missing_files():
    """Test that path resolver returns None for files that can't be resolved"""
    
    with tempfile.TemporaryDirectory() as temp_dir:
        uploads_dir = os.path.join(temp_dir, 'uploads')
        os.makedirs(uploads_dir)
        
        with patch('app.utils.path_resolver._get_uploads_directory', return_value=uploads_dir):
            # Test nonexistent file
            windows_path = r'D:\Github\evzj\uploads\nonexistent.jpg'
            resolved_path = resolve_upload_path(windows_path)
            
            assert resolved_path is None


def test_friendly_message_is_localized():
    """Test that friendly message is properly localized"""
    message = get_friendly_image_message()
    assert message == "图片缺失或不可访问"
    assert isinstance(message, str)
    assert len(message) > 0


def test_template_fallback_uses_friendly_message():
    """Test that template fallback logic uses friendly messages instead of raw paths"""
    from app.reporting.docx_renderer import _create_minimal_template
    import tempfile
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        try:
            _create_minimal_template(tmp_file.name)
            
            # Read the template content
            from docx import Document
            doc = Document(tmp_file.name)
            
            # Extract template content as text
            template_text = ""
            for paragraph in doc.paragraphs:
                template_text += paragraph.text + " "
            
            # Ensure no raw path fallbacks in template (these are the problematic fallbacks)
            assert "}}{{ images.composited_image_path }}" not in template_text
            assert "}}{{ images.original_image_path }}" not in template_text
            # Should have friendly message fallback
            assert "图片缺失或不可访问" in template_text or "friendly_message" in template_text
            
        finally:
            os.unlink(tmp_file.name)


def test_assignment_template_fallback_uses_friendly_message():
    """Test that assignment template fallback logic uses friendly messages"""
    from app.reporting.docx_renderer import _create_assignment_template
    import tempfile
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        try:
            _create_assignment_template(tmp_file.name)
            
            # Read the template content
            from docx import Document
            doc = Document(tmp_file.name)
            
            # Extract template content as text
            template_text = ""
            for paragraph in doc.paragraphs:
                template_text += paragraph.text + " "
            
            # Ensure no raw path fallbacks in template (these are the problematic fallbacks)
            assert "}}{{ s.images.composited_image_path }}" not in template_text
            assert "}}{{ s.images.original_image_path }}" not in template_text
            # Should have friendly message fallback
            assert "图片缺失或不可访问" in template_text or "friendly_message" in template_text
            
        finally:
            os.unlink(tmp_file.name)


def test_service_imports_path_resolver():
    """Test that service module correctly imports path resolver functions"""
    # Test that the functions are available in the module scope by importing locally
    try:
        from app.utils.path_resolver import resolve_upload_path, get_friendly_image_message
        # If we can import these, the integration is working
        assert callable(resolve_upload_path)
        assert callable(get_friendly_image_message)
        # Service module should be able to import them too
        import app.reporting.service
        # This will raise ImportError if service.py can't import the functions
    except ImportError as e:
        pytest.fail(f"Path resolver functions should be importable: {e}")


if __name__ == '__main__':
    pytest.main([__file__])