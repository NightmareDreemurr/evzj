"""
Tests for improved image error handling in DOCX reports.

This test specifically addresses the issue where missing images showed "None"
instead of user-friendly error messages.
"""
import tempfile
import pytest
from app.schemas.evaluation import EvaluationResult, Meta, Scores, RubricScore, TextBlock
from app.reporting.docx_renderer import render_essay_docx


class MockEssayWithImages:
    """Mock essay instance with image paths for testing"""
    def __init__(self, original_path=None, overlay_path=None):
        self.original_image_path = original_path
        self.annotated_overlay_path = overlay_path


def test_missing_images_show_friendly_message():
    """Test that missing images show friendly error message instead of 'None'"""
    
    # Create evaluation with non-existent image paths
    evaluation = EvaluationResult(
        meta=Meta(
            student="测试学生",
            class_="测试班级", 
            teacher="测试老师",
            topic="测试作文",
            date="2024-08-21"
        ),
        scores=Scores(
            total=80.0,
            rubrics=[RubricScore(name="内容", score=16.0, max=20.0, weight=1.0, reason="测试")]
        ),
        text=TextBlock(original="测试内容", cleaned="测试内容")
    )
    
    # Attach mock essay with invalid image paths (simulating the reported issue)
    evaluation._essay_instance = MockEssayWithImages(
        original_path="D:\\Github\\evzj\\uploads\\nonexistent.jpg",
        overlay_path="D:\\Github\\evzj\\uploads\\nonexistent_overlay.jpg"
    )
    
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation)
        
        # Read the generated content
        from docx import Document
        doc = Document(result_path)
        
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Should NOT contain "None"
        assert "None" not in full_text, "DOCX should not contain 'None' for missing images"
        
        # Should contain friendly error message
        assert "图片缺失或不可访问" in full_text, "DOCX should contain friendly error message for missing images"
        
        # Should still have the image section header
        assert "作文图片" in full_text, "DOCX should still have image section header"


def test_no_images_no_error_message():
    """Test that evaluations without image paths don't show error messages"""
    
    evaluation = EvaluationResult(
        meta=Meta(student="测试学生", topic="测试作文", date="2024-08-21"),
        scores=Scores(total=80.0, rubrics=[]),
        text=TextBlock(original="测试内容", cleaned="测试内容")
    )
    
    # No _essay_instance means no images
    
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation)
        
        from docx import Document
        doc = Document(result_path)
        
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Should not contain error messages when no images are expected
        assert "图片缺失或不可访问" not in full_text
        assert "None" not in full_text


def test_empty_image_paths_handled():
    """Test that empty string image paths are handled gracefully"""
    
    evaluation = EvaluationResult(
        meta=Meta(student="测试学生", topic="测试作文", date="2024-08-21"),
        scores=Scores(total=80.0, rubrics=[]),
        text=TextBlock(original="测试内容", cleaned="测试内容")
    )
    
    # Essay with empty string paths
    evaluation._essay_instance = MockEssayWithImages(
        original_path="",
        overlay_path=""
    )
    
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation)
        
        from docx import Document
        doc = Document(result_path)
        
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Empty paths should be treated as no images
        assert "None" not in full_text
        assert "图片缺失或不可访问" not in full_text


def test_single_invalid_path_handled():
    """Test that having only one invalid image path is handled correctly"""
    
    evaluation = EvaluationResult(
        meta=Meta(student="测试学生", topic="测试作文", date="2024-08-21"),
        scores=Scores(total=80.0, rubrics=[]),
        text=TextBlock(original="测试内容", cleaned="测试内容")
    )
    
    # Only original image path, invalid
    evaluation._essay_instance = MockEssayWithImages(
        original_path="/invalid/path/image.jpg",
        overlay_path=None
    )
    
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation)
        
        from docx import Document
        doc = Document(result_path)
        
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Should show friendly error message for the invalid path
        assert "None" not in full_text
        assert "图片缺失或不可访问" in full_text
        assert "作文图片" in full_text