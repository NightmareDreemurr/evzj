#!/usr/bin/env python3
"""
Test case for the image rendering fix in DOCX generation.
"""
import os
import tempfile
import pytest

from app.schemas.evaluation import EvaluationResult, Meta, Scores, TextBlock
from app.reporting.docx_renderer import render_essay_docx


class TestImageRenderingFix:
    """Test the image rendering fix for DOCX generation"""
    
    def create_test_image(self):
        """Create a temporary test image file"""
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as f:
            # Create a minimal JPEG header to make it a valid image file
            f.write(b'\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x01\x00H\x00H\x00\x00\xff\xdb\x00C\x00\x08\x06\x06\x07\x06\x05\x08\x07\x07\x07\t\t\x08\n\x0c\x14\r\x0c\x0b\x0b\x0c\x19\x12\x13\x0f\x14\x1d\x1a\x1f\x1e\x1d\x1a\x1c\x1c $.\' ",#\x1c\x1c(7),01444\x1f\'9=82<.342\xff\xc0\x00\x11\x08\x00\x01\x00\x01\x01\x01\x11\x00\x02\x11\x01\x03\x11\x01\xff\xc4\x00\x14\x00\x01\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x08\xff\xc4\x00\x14\x10\x01\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\xff\xda\x00\x08\x01\x01\x00\x00?\x00\xff\xd9')
            temp_path = f.name
        return temp_path
    
    def test_image_rendering_without_missing_message(self):
        """Test that images are rendered without showing the '图片缺失或不可访问' message"""
        # Create a temporary test image
        test_image_path = self.create_test_image()
        
        try:
            # Create test evaluation data
            meta = Meta(
                student="测试学生",
                class_="测试班级",
                teacher="测试教师",
                topic="测试作文",
                date="2024-01-01",
                student_id="123",
                grade="五年级",
                words=100
            )
            
            scores = Scores(total=85.0, rubrics=[])
            text = TextBlock(original="测试原文内容", cleaned="测试清洗后内容")
            
            evaluation = EvaluationResult(meta=meta, scores=scores, text=text)
            
            # Mock an essay instance with our test image
            class MockEssay:
                def __init__(self):
                    self.original_image_path = test_image_path
                    self.annotated_overlay_path = None
            
            # Attach the mock essay to the evaluation
            evaluation._essay_instance = MockEssay()
            
            # Generate DOCX
            with tempfile.TemporaryDirectory() as temp_dir:
                output_path = os.path.join(temp_dir, "test_image_rendering.docx")
                result_path = render_essay_docx(evaluation, output_path)
                
                # Verify the file was created
                assert os.path.exists(result_path)
                assert os.path.getsize(result_path) > 0
                
                # Extract text from the generated DOCX and check for missing image message
                from docx import Document
                doc = Document(result_path)
                full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                # The main assertion - no missing image message should appear
                assert "图片缺失或不可访问" not in full_text, \
                    "Document should not contain missing image message when valid image is provided"
                
                # Should have an image section
                assert "作文图片" in full_text, \
                    "Document should contain image section when image is available"
                    
        finally:
            # Clean up the temporary image
            if os.path.exists(test_image_path):
                os.unlink(test_image_path)
    
    def test_missing_image_fallback_message(self):
        """Test that missing images show the appropriate fallback message"""
        # Create test evaluation data without any image
        meta = Meta(
            student="测试学生",
            class_="测试班级", 
            teacher="测试教师",
            topic="测试作文",
            date="2024-01-01",
            student_id="123",
            grade="五年级",
            words=100
        )
        
        scores = Scores(total=85.0, rubrics=[])
        text = TextBlock(original="测试原文内容", cleaned="测试清洗后内容")
        
        evaluation = EvaluationResult(meta=meta, scores=scores, text=text)
        
        # Mock an essay instance with a non-existent image path
        class MockEssay:
            def __init__(self):
                self.original_image_path = "/nonexistent/path/image.jpg"
                self.annotated_overlay_path = None
        
        # Attach the mock essay to the evaluation
        evaluation._essay_instance = MockEssay()
        
        # Generate DOCX
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = os.path.join(temp_dir, "test_missing_image.docx")
            result_path = render_essay_docx(evaluation, output_path)
            
            # Verify the file was created
            assert os.path.exists(result_path)
            assert os.path.getsize(result_path) > 0
            
            # Extract text from the generated DOCX
            from docx import Document
            doc = Document(result_path)
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Should contain the friendly message when image cannot be found
            assert "图片缺失或不可访问" in full_text, \
                "Document should contain missing image message when image path cannot be resolved"
    
    def test_no_image_essay(self):
        """Test that essays without any image information don't show image sections"""
        # Create test evaluation data
        meta = Meta(
            student="测试学生",
            class_="测试班级",
            teacher="测试教师", 
            topic="测试作文",
            date="2024-01-01",
            student_id="123",
            grade="五年级",
            words=100
        )
        
        scores = Scores(total=85.0, rubrics=[])
        text = TextBlock(original="测试原文内容", cleaned="测试清洗后内容")
        
        evaluation = EvaluationResult(meta=meta, scores=scores, text=text)
        
        # Mock an essay instance with no image paths
        class MockEssay:
            def __init__(self):
                self.original_image_path = None
                self.annotated_overlay_path = None
        
        # Attach the mock essay to the evaluation  
        evaluation._essay_instance = MockEssay()
        
        # Generate DOCX
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = os.path.join(temp_dir, "test_no_image.docx")
            result_path = render_essay_docx(evaluation, output_path)
            
            # Verify the file was created
            assert os.path.exists(result_path)
            assert os.path.getsize(result_path) > 0
            
            # Extract text from the generated DOCX
            from docx import Document
            doc = Document(result_path)
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Should not contain image section or missing image message
            assert "作文图片" not in full_text, \
                "Document should not contain image section when no image paths are provided"
            assert "图片缺失或不可访问" not in full_text, \
                "Document should not contain missing image message when no image paths are provided"