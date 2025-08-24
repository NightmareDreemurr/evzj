"""
Tests to verify that real AI evaluation data is used instead of hardcoded fallbacks.

This test specifically addresses the issue where DOCX renderer was using
fake/hardcoded data instead of reading from the database.
"""
import tempfile
import pytest
from app.schemas.evaluation import EvaluationResult, Meta, Scores, RubricScore, TextBlock
from app.reporting.docx_renderer import render_essay_docx


def test_real_ai_data_is_used_instead_of_fallbacks():
    """Test that real AI evaluation data is used instead of hardcoded fallback text"""
    
    # Create evaluation with real AI data structure similar to database sample
    evaluation = EvaluationResult(
        meta=Meta(
            student="蔡桢达",
            class_="",
            teacher="沈颖",
            topic="写读后感--实用-学术类设计",
            date="2025-08-23",
            student_id="47",
            grade="",
            words=0
        ),
        scores=Scores(
            total=32,
            rubrics=[
                RubricScore(
                    name="书籍内容概述提炼",
                    score=17,
                    max=20,
                    weight=1.0,
                    reason="在这一维度上表现优秀！",
                    level="A",
                    example_good_sentence=["《魔道祖师》是墨香铜臭所著的长篇修真小说，主要讲述了夷陵老祖魏无羡献舍归来后，与姑苏蓝氏二公子蓝忘机共同探查莫家庄凶尸奇案的故事。"],
                    example_improvement_suggestion=[]
                ),
                RubricScore(
                    name="心得体会交流分享",
                    score=8,
                    max=10,
                    weight=1.0,
                    reason="心得体会部分做得很好！",
                    level="A",
                    example_good_sentence=["记得有一次，我和一个朋友开了一个玩笑，说他有点矮，结果出去玩时他拿着篮球砸向我的头。"],
                    example_improvement_suggestion=[]
                ),
                RubricScore(
                    name="语言表达学术规范",
                    score=7,
                    max=10,
                    weight=1.0,
                    reason="语言表达基本规范，但还有一些提升空间。",
                    level="B",
                    example_good_sentence=["《魔道祖师》是墨香铜臭所著的长篇修真小说"],
                    example_improvement_suggestion=["优秀的文学作品往往能够给予读者深刻的人生启示。"]
                )
            ]
        ),
        overall_comment="这篇读后感展现了不错的阅读理解和感悟能力！你能够清晰地概述《魔道祖师》的主要内容和人物特点，特别是对观音庙情节的描写很具体，并且能够从中提炼出'祸从口出'的道理，还能结合自己的生活经历来谈体会，这一点非常值得肯定。整体结构完整，思路清晰，是一篇有思考、有感悟的读后感。",
        strengths=[
            "能够准确概括书籍主要内容并聚焦自己感兴趣的情节进行详细描述",
            "能够将阅读感悟与自身生活经历相结合，体现了真实的阅读收获"
        ],
        improvements=[
            "可以进一步深化对书中主题和人物命运的理解，挖掘更深层的启示",
            "语言表达可以更加学术化和规范化，减少口语化表述"
        ],
        text=TextBlock(
            original="读《魔道祖师》有感",
            cleaned="读《魔道祖师》有感"
        ),
        assignmentTitle="读后感写作练习",
        studentName="蔡桢达",
        submittedAt="2025-08-23",
        currentEssayContent="读《魔道祖师》有感\n最近，我读了作者墨香铜臭的《魔道祖师》，颇有感触。"
    )
    
    # Generate DOCX
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation)
        
        # Read the generated content
        from docx import Document
        doc = Document(result_path)
        
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Verify real strengths are used instead of fallbacks
        real_strength = "能够准确概括书籍主要内容并聚焦自己感兴趣的情节进行详细描述"
        fallback_strength = "能够完成作文基本要求"
        assert real_strength in full_text, "Real strength data should be used"
        assert fallback_strength not in full_text, "Fallback strength should not be used when real data exists"
        
        # Verify real improvements are used instead of fallbacks
        real_improvement = "可以进一步深化对书中主题和人物命运的理解，挖掘更深层的启示"
        fallback_improvement = "建议进一步丰富表达方式，提升语言准确性"
        assert real_improvement in full_text, "Real improvement data should be used"
        assert fallback_improvement not in full_text, "Fallback improvement should not be used when real data exists"
        
        # Verify real overall comment is used instead of fallbacks
        real_overall_comment = "这篇读后感展现了不错的阅读理解和感悟能力"
        fallback_overall_comment = "本次作文总体表现良好"
        assert real_overall_comment in full_text, "Real overall comment should be used"
        assert fallback_overall_comment not in full_text, "Fallback overall comment should not be used when real data exists"
        
        # Verify real example sentences are used instead of fallbacks
        real_example_sentence = "《魔道祖师》是墨香铜臭所著的长篇修真小说"
        fallback_example_sentence = "文章基本符合要求，表达较为清楚"
        assert real_example_sentence in full_text, "Real example sentence should be used"
        assert fallback_example_sentence not in full_text, "Fallback example sentence should not be used when real data exists"


def test_fallbacks_used_when_data_missing():
    """Test that fallbacks are appropriately used when real data is missing"""
    
    # Create evaluation with minimal data (no AI enhancement data)
    evaluation = EvaluationResult(
        meta=Meta(
            student="测试学生",
            class_="测试班级",
            teacher="测试老师",
            topic="测试作文",
            date="2024-08-21"
        ),
        scores=Scores(
            total=90.0,
            rubrics=[
                RubricScore(name="测试维度", score=18.0, max=20.0, weight=1.0, reason="测试理由")
            ]
        ),
        text=TextBlock(
            original="这是测试文本内容。",
            cleaned="这是测试文本内容。"
        ),
        # Explicitly set empty data to simulate missing AI enhancement
        strengths=[],
        improvements=[],
        overall_comment=""
    )
    
    # Generate DOCX
    with tempfile.TemporaryDirectory() as temp_dir:
        result_path = render_essay_docx(evaluation)
        
        # Read the generated content
        from docx import Document
        doc = Document(result_path)
        
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # When data is missing, fallbacks should be used
        fallback_strength = "能够完成作文基本要求"
        fallback_improvement = "可以进一步丰富内容深度"
        fallback_example_sentence = "文章基本符合要求，表达较为清楚"
        
        assert fallback_strength in full_text, "Fallback strength should be used when data is missing"
        assert fallback_improvement in full_text, "Fallback improvement should be used when data is missing"
        assert fallback_example_sentence in full_text, "Fallback example sentence should be used when data is missing"